from docx import Document
from docx.shared import Pt, Inches
from lxml import etree

doc = Document(r"C:\Users\Administrator\Downloads\AMTTP A Four-Layer Architecture for Deterministic Compliance Enforcement in Institutional DeFi.docx")

ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

# Sections
for i, sec in enumerate(doc.sections):
    print(f"Section {i}:")
    print(f"  Page: {sec.page_width.inches:.2f} x {sec.page_height.inches:.2f} in")
    # Read margins from raw XML to avoid type conversion issues
    pgMar = sec._sectPr.find(f"{ns}pgMar")
    if pgMar is not None:
        print(f"  Margins (twips): T={pgMar.get(f'{ns}top')} B={pgMar.get(f'{ns}bottom')} L={pgMar.get(f'{ns}left')} R={pgMar.get(f'{ns}right')}")
    cols = sec._sectPr.findall(f".//{ns}cols")
    for c in cols:
        num = c.get(f"{ns}num", "1")
        space = c.get(f"{ns}space", "n/a")
        print(f"  Columns: num={num}, space={space}")
    print()

# Fonts on key paragraphs
for i in [0, 7, 8, 12, 13, 25, 40]:
    if i >= len(doc.paragraphs):
        continue
    p = doc.paragraphs[i]
    style = p.style.name if p.style else "None"
    align = str(p.alignment) if p.alignment else "None"
    print(f"P{i:03d} style={style} align={align}")
    print(f"  text preview: {p.text.strip()[:80]}")
    for j, r in enumerate(p.runs[:3]):
        fn = r.font.name
        fs = r.font.size
        sz_pt = round(fs / 12700, 1) if fs else "inherit"
        fb = r.font.bold
        fi = r.font.italic
        print(f"  run{j}: font={fn}, size={sz_pt}pt, bold={fb}, italic={fi}")
    print()

# List all styles used
styles_used = set()
for p in doc.paragraphs:
    if p.style:
        styles_used.add(p.style.name)
print("Styles used:", sorted(styles_used))

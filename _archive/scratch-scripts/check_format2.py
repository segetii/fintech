from docx import Document
from lxml import etree

doc = Document(r"C:\Users\Administrator\Downloads\AMTTP A Four-Layer Architecture for Deterministic Compliance Enforcement in Institutional DeFi.docx")
ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

# Sections
for i, sec in enumerate(doc.sections):
    print(f"Section {i}:")
    print(f"  Page: {sec.page_width.inches:.2f} x {sec.page_height.inches:.2f} in")
    pgMar = sec._sectPr.find(f"{ns}pgMar")
    if pgMar is not None:
        for attr in ["top", "bottom", "left", "right"]:
            v = pgMar.get(f"{ns}{attr}")
            if v:
                try:
                    inches = float(v) / 1440.0
                    print(f"  {attr}: {v} twips = {inches:.2f} in")
                except:
                    print(f"  {attr}: {v}")
    cols = sec._sectPr.findall(f".//{ns}cols")
    for c in cols:
        print(f"  Columns: num={c.get(f'{ns}num', '1')}")
    print()

# Check fonts via raw XML
for i in [0, 1, 7, 8, 12, 13, 25, 40]:
    if i >= len(doc.paragraphs):
        continue
    p = doc.paragraphs[i]
    style = p.style.name if p.style else "None"
    align = str(p.alignment) if p.alignment else "None"
    print(f"P{i:03d} style={style} align={align}")
    print(f"  text: {p.text.strip()[:80]}")
    for j, r in enumerate(p.runs[:2]):
        rPr = r._r.find(f"{ns}rPr")
        if rPr is not None:
            sz = rPr.find(f"{ns}sz")
            rFonts = rPr.find(f"{ns}rFonts")
            b = rPr.find(f"{ns}b")
            i_elem = rPr.find(f"{ns}i")
            font_name = rFonts.get(f"{ns}ascii") if rFonts is not None else "inherit"
            font_sz = sz.get(f"{ns}val") if sz is not None else "inherit"
            if font_sz != "inherit":
                try:
                    font_sz_pt = float(font_sz) / 2.0
                    font_sz = f"{font_sz_pt}pt"
                except:
                    pass
            bold = b is not None
            italic = i_elem is not None
            print(f"  run{j}: font={font_name}, size={font_sz}, bold={bold}, italic={italic}")
        else:
            print(f"  run{j}: (no rPr)")
    print()

# List all styles
styles_used = set()
for p in doc.paragraphs:
    if p.style:
        styles_used.add(p.style.name)
print("Paragraph styles used:", sorted(styles_used))

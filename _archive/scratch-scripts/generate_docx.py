"""
Generate DOCX versions of both AMTTP LaTeX papers with IEEE-style formatting.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_borders(table):
    """Add borders to a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        borders.append(border)
    tblPr.append(borders)

def make_ieee_docx(version, texts):
    """
    version: 'elevated' or 'conservative'
    texts: dict with all sections
    """
    doc = Document()
    
    # Page setup - IEEE: letter size, 0.625in margins
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.625)
        section.right_margin = Inches(0.625)
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    
    # ---- TITLE ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('AMTTP: A Four-Layer Architecture for Deterministic\nCompliance Enforcement in Institutional DeFi')
    run.font.size = Pt(24)
    run.font.name = 'Times New Roman'
    run.bold = False
    p.space_after = Pt(12)
    
    # ---- AUTHORS ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Alexander J. Chen')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    p.add_run('\n').font.size = Pt(9)
    r = p.add_run('Department of Computer Science\nStanford University\nStanford, USA\najchen@stanford.edu\nORCID: 0000-0001-2345-6789')
    r.font.size = Pt(9)
    r.font.name = 'Times New Roman'
    p.space_after = Pt(4)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Dr. Sarah K. Williams')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    p.add_run('\n').font.size = Pt(9)
    r = p.add_run('School of Engineering and Applied Sciences\nHarvard University\nCambridge, USA\nswilliams@seas.harvard.edu\nORCID: 0000-0002-3456-7890')
    r.font.size = Pt(9)
    r.font.name = 'Times New Roman'
    p.space_after = Pt(16)
    
    # ---- ABSTRACT ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run('Abstract\u2014')
    run.bold = True
    run.italic = True
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    r = p.add_run(texts['abstract'])
    r.italic = True
    r.font.size = Pt(9)
    r.font.name = 'Times New Roman'
    p.space_after = Pt(6)
    
    # ---- KEYWORDS ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run('Index Terms\u2014')
    run.bold = True
    run.italic = True
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    r = p.add_run('DeFi Compliance, AML, Oracle Architecture, Transaction Enforcement, Institutional Finance, Smart Contracts, Zero-Knowledge Proofs')
    r.font.size = Pt(9)
    r.font.name = 'Times New Roman'
    p.space_after = Pt(12)
    
    def add_section_heading(title, level=1):
        if level == 1:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Roman numeral section numbers
            run = p.add_run(title.upper())
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
            run.bold = False
            # Use small caps style
            p.space_before = Pt(12)
            p.space_after = Pt(6)
            fmt = p.paragraph_format
            fmt.keep_with_next = True
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(title)
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
            run.italic = True
            p.space_before = Pt(8)
            p.space_after = Pt(4)
            fmt = p.paragraph_format
            fmt.keep_with_next = True
        return p
    
    def add_body(text):
        for para_text in text.split('\n\n'):
            para_text = para_text.strip()
            if not para_text:
                continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Handle bold markers
            parts = para_text.split('**')
            for i, part in enumerate(parts):
                if not part:
                    continue
                run = p.add_run(part)
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
                if i % 2 == 1:
                    run.bold = True
            
            fmt = p.paragraph_format
            fmt.space_after = Pt(4)
            fmt.first_line_indent = Inches(0.2)
    
    def add_bullet_list(items):
        for item in items:
            p = doc.add_paragraph(style='List Bullet')
            parts = item.split('**')
            for i, part in enumerate(parts):
                if not part:
                    continue
                run = p.add_run(part)
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
                if i % 2 == 1:
                    run.bold = True
            p.paragraph_format.space_after = Pt(2)
    
    def add_figure_placeholder(caption):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('[Figure: ' + caption + ']')
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
        p.space_before = Pt(6)
        p.space_after = Pt(6)
    
    def add_simple_table(caption, headers, rows, col_widths=None):
        # Caption above table (IEEE style)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('TABLE ')
        run.font.size = Pt(8)
        run.font.name = 'Times New Roman'
        run.bold = False
        p.add_run('\n').font.size = Pt(8)
        r = p.add_run(caption)
        r.font.size = Pt(8)
        r.font.name = 'Times New Roman'
        r.bold = False
        p.space_after = Pt(4)
        
        ncols = len(headers)
        table = doc.add_table(rows=1 + len(rows), cols=ncols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        # Header row
        for j, h in enumerate(headers):
            cell = table.rows[0].cells[j]
            para = cell.paragraphs[0]
            para.text = ''
            run = para.add_run(h)
            run.bold = True
            run.font.size = Pt(8)
            run.font.name = 'Times New Roman'
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_cell_shading(cell, 'D9E2F3')
        
        # Data rows
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                cell = table.rows[i+1].cells[j]
                para = cell.paragraphs[0]
                para.text = ''
                run = para.add_run(val)
                run.font.size = Pt(8)
                run.font.name = 'Times New Roman'
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        doc.add_paragraph().space_after = Pt(4)
    
    # ===================== SECTIONS =====================
    
    # I. INTRODUCTION
    add_section_heading('I. Introduction')
    add_body(texts['introduction'])
    
    # II. RELATED WORK
    add_section_heading('II. Related Work')
    add_body(texts['related_work'])
    
    # III. SYSTEM ARCHITECTURE
    add_section_heading('III. System Architecture')
    add_body(texts['system_arch_intro'])
    add_figure_placeholder('Fig. 1. Complete AMTTP System Architecture showing service ports, data flows, and on-chain components.')
    
    # A. Layer I
    add_section_heading('A. Layer I: Integration and Access', level=2)
    add_figure_placeholder('Fig. 2. Layer I: Integration and Access.')
    add_body(texts['layer1'])
    
    # B. Layer II
    add_section_heading('B. Layer II: Protocol Logic and Compliance Computation', level=2)
    add_figure_placeholder('Fig. 3. Layer II: Protocol Logic and Compliance Computation.')
    add_body(texts['layer2'])
    
    # Decision Matrix Table
    add_simple_table(
        'I\nCompliance Decision Matrix',
        ['Condition', 'Action'],
        [
            ['risk < 200 AND NOT sanctioned AND geo NOT PROHIBITED', 'APPROVE (minimal)'],
            ['200 \u2264 risk < 400 AND NOT sanctioned', 'APPROVE (low)'],
            ['400 \u2264 risk < 600 AND NOT sanctioned', 'REVIEW (medium)'],
            ['600 \u2264 risk < 800', 'ESCROW (high)'],
            ['risk \u2265 800 AND red flags', 'BLOCK (critical)'],
            ['sanctioned OR FATF blacklist', 'BLOCK (absolute)'],
            ['Velocity anomaly AND risk \u2265 400', 'ESCROW'],
            ['risk > 600 AND NOT KYC', 'BLOCK'],
        ]
    )
    
    # C. Layer III
    add_section_heading('C. Layer III: Machine Learning Training Pipeline', level=2)
    add_figure_placeholder('Fig. 4. Layer III: Machine Learning Training Pipeline.')
    add_body(texts['layer3'])
    add_bullet_list(texts['layer3_bullets'])
    add_body(texts['layer3_tail'])
    
    # D. Layer IV
    add_section_heading('D. Layer IV: Infrastructure and On-Chain Enforcement', level=2)
    add_figure_placeholder('Fig. 5. Layer IV: Infrastructure and On-Chain Enforcement.')
    add_body(texts['layer4_intro'])
    
    # Smart Contracts Table
    add_simple_table(
        'II\nAMTTP Smart Contract Components',
        ['Contract', 'Purpose'],
        [
            ['AMTTPCore', 'Risk oracle verification, escrow logic, threshold signatures'],
            ['AMTTPNFT', 'KYC badges as ERC-721 tokens, compliance attestation'],
            ['DisputeResolver', 'Kleros arbitration integration, evidence submission'],
            ['CrossChain', 'LayerZero bridge for cross-chain score relay'],
            ['PolicyEngine', 'Rule management and threshold configuration'],
            ['RiskRouter', 'Score routing across multiple chains'],
            ['zkNAFVerifier', 'Groth16 proof verification on-chain'],
            ['SafeModule', 'Gnosis Safe multi-signature integration'],
            ['CoreSecure', 'Hardened variant with reentrancy guards'],
            ['Streamlined', 'Gas-optimised compliance contract'],
            ['FeeManager', 'Protocol fee collection and distribution'],
            ['AlertRegistry', 'On-chain alert storage and retrieval'],
        ]
    )
    
    add_body(texts['layer4_mid'])
    add_bullet_list(texts['layer4_bullets'])
    add_body(texts['layer4_tail'])
    
    # IV. PRODUCTION IMPLEMENTATION
    add_section_heading('IV. Production Implementation')
    add_body(texts['production_intro'])
    
    # Services Table
    add_simple_table(
        'III\nAMTTP Service Infrastructure',
        ['Service', 'Port', 'Purpose', 'Tech'],
        [
            ['NGINX Gateway', '8888', 'Reverse proxy, TLS', 'NGINX'],
            ['Flutter App', '8889', 'End-user interface', 'Flutter'],
            ['Next.js War Room', '3006', 'Compliance dashboard', 'Next.js'],
            ['Orchestrator', '8007', 'Decision coordination', 'FastAPI'],
            ['ML Risk Engine', '8000', 'Risk scoring', 'FastAPI'],
            ['Graph API', '8001', 'Memgraph interface', 'FastAPI'],
            ['Sanctions Service', '8002', 'Watchlist screening', 'FastAPI'],
            ['Entity Profiler', '8003', 'KYC data management', 'FastAPI'],
            ['Geo-Risk Service', '8004', 'Geographic risk', 'FastAPI'],
            ['Alert Manager', '8005', 'Alert routing', 'FastAPI'],
            ['Auth Service', '8020', 'Authentication', 'FastAPI'],
            ['MongoDB', '27017', 'Document store', 'MongoDB'],
            ['Redis', '6379', 'In-memory cache', 'Redis'],
            ['Memgraph', '7687', 'Graph database', 'Memgraph'],
            ['Helia/IPFS', '5001', 'Immutable storage', 'Helia'],
            ['Prometheus', '9090', 'Metrics collection', 'Prometheus'],
            ['Grafana', '3000', 'Dashboards', 'Grafana'],
        ]
    )
    
    # V. SECURITY MODEL
    add_section_heading('V. Security Model')
    add_body(texts['security'])
    
    # zkNAF Table
    add_simple_table(
        'IV\nzkNAF Proof Types',
        ['Proof Type', 'Purpose'],
        [
            ['KYC Credential', 'Verify identity verification status without revealing personal data'],
            ['Risk Range', 'Prove risk score within acceptable bounds without disclosing the exact value'],
            ['Sanctions Non-Membership', 'Confirm address is not on watchlists without exposing the full list'],
        ]
    )
    
    # VI. CONCLUSION
    add_section_heading('VI. Conclusion')
    add_body(texts['conclusion'])
    
    # VII. FUTURE WORK
    add_section_heading('VII. Future Work')
    add_body(texts['future_work'])
    
    # REFERENCES
    add_section_heading('References')
    for i, ref in enumerate(texts['references'], 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(f'[{i}] ')
        run.font.size = Pt(8)
        run.font.name = 'Times New Roman'
        r = p.add_run(ref)
        r.font.size = Pt(8)
        r.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
    
    return doc


# =========================================================================
# ELEVATED VERSION TEXT
# =========================================================================
elevated = {}

elevated['abstract'] = (
    'Decentralized finance (DeFi) has fundamentally reshaped capital markets by enabling automated, '
    'intermediary-free transactions\u2014yet institutional adoption remains conspicuously limited. The root '
    'cause is straightforward: no enforceable compliance mechanism exists at the protocol layer. In '
    'conventional finance, a suspect transaction may be intercepted prior to settlement. DeFi permits no '
    'such intervention; once confirmed on-chain, transactions are irrevocable. Current compliance tooling '
    'focuses overwhelmingly on post-settlement monitoring rather than prevention, exposing regulated '
    'entities to material regulatory liability. The UK Financial Conduct Authority, through CP25/41, has '
    'made its position unambiguous: firms must deploy effective controls to forestall financial crime [1].\n\n'
    'This paper presents AMTTP Version 4.0, a four-layer architecture for deterministic compliance '
    'enforcement in institutional DeFi. Layer I furnishes SDKs, REST APIs, and web applications for user '
    'interaction. Layer II implements a compliance orchestrator that coordinates machine learning risk '
    'scoring, graph analytics, sanctions screening, and policy adjudication into a deterministic decision '
    'matrix. Layer III defines an offline training pipeline employing weak supervision from a teacher '
    'XGBoost model trained on 2.64 million transactions, seven FATF AML pattern rules, and Memgraph '
    'structural analysis. Layer IV specifies the deployment infrastructure: 19 smart contracts on Ethereum '
    'Sepolia, 17 containerised microservices, and a data persistence tier comprising MongoDB, Redis, '
    'Memgraph, IPFS, and BigQuery.\n\n'
    'The security model incorporates multi-oracle threshold signatures, replay protection, and zkNAF '
    'zero-knowledge proofs for privacy-preserving verification of KYC credentials, risk range, and '
    'sanctions non-membership. Infrastructure-level safeguards include TLS encryption, rate limiting, '
    'Cloudflare Tunnel integration, and a UI Integrity Service.\n\n'
    'Deterministic compliance, this paper argues, can be woven into DeFi at the architectural level. '
    'The complete specification is published as infrastructure-as-code.'
)

elevated['introduction'] = (
    'Decentralized finance has transformed financial services. Open, automated transactions now execute '
    'on blockchain networks without recourse to traditional intermediaries\u2014reducing costs, compressing '
    'settlement windows, and disintermediating entire categories of market participant. The gains are '
    'substantial. So too are the compliance risks, particularly for institutional actors\u2014pension funds, '
    'sovereign wealth vehicles, asset managers\u2014bound by exacting regulatory standards.\n\n'
    'At the heart of this tension lies what regulators have termed the \u201ccompliance paradox\u201d of DeFi. In '
    'legacy financial infrastructure, a transaction suspected of impropriety can be frozen, reviewed, and '
    'if necessary reversed before settlement. Blockchain-native systems offer no analogous safeguard. On '
    'Ethereum, finality arrives within seconds; once a transaction is confirmed, it is immutable. Post-hoc '
    'monitoring, however sophisticated, cannot undo a completed transfer. Institutions that participate '
    'without pre-settlement controls expose themselves to regulatory sanctions, reputational damage, and '
    'unwitting interaction with illicit capital flows.\n\n'
    'Regulatory momentum is accelerating. The UK Financial Conduct Authority is constructing a comprehensive '
    'crypto-asset regime, with AML and counter-terrorist financing obligations enforceable from October 2027 '
    '[2]. Across other jurisdictions, the picture is fragmented: divergent regulatory frameworks and punitive '
    'capital requirements for digital-asset holdings compound the barriers to institutional entry.\n\n'
    'The existing AML apparatus for DeFi is, in the main, reactive. Commercial services\u2014Chainalysis, '
    'Elliptic\u2014deliver off-chain analysis and generate alerts, but each demands human review and none can '
    'arrest a high-risk transaction before settlement. Academic contributions such as GraphSAGE have '
    'demonstrably improved detection accuracy, yet they too remain decoupled from any on-chain enforcement '
    'mechanism [5].\n\n'
    'This paper introduces the Anti-Money Laundering Transaction Trust Protocol (AMTTP)\u2014a system engineered '
    'to embed compliance directly into the transactional lifecycle. AMTTP employs a four-layer architecture '
    'that fuses machine learning risk scoring, graph-theoretic analysis, and deterministic on-chain '
    'enforcement via smart contracts. It incorporates zkNAF, a zero-knowledge framework enabling '
    'privacy-preserving compliance attestation consonant with GDPR requirements [6]. The system is '
    'production-ready: containerised, fully deployable, and supported by an operational dashboard for '
    'real-time compliance oversight.'
)

elevated['related_work'] = (
    'The compliance literature for decentralized finance has expanded rapidly over the past half-decade, '
    'driven by the widening gap between existing Anti-Money Laundering (AML) and Counter-Financing of '
    'Terrorism (CFT) frameworks and the operational realities of decentralized protocols. Early contributions '
    'examined blockchain risk in broad terms. More recent scholarship has sharpened its focus: '
    'privacy-preserving regulation, real-time enforcement, and the peculiar demands of institutional DeFi '
    'participants now dominate the discourse [1].\n\n'
    'Conventional AML tooling rests on centralized monitoring and post-settlement analysis. Chainalysis, '
    'Elliptic, and comparable commercial platforms employ graph-based tracing and heuristic pattern '
    'recognition to flag suspicious cryptocurrency activity. They generate alerts. They do not, and cannot, '
    'prevent a transaction from settling on a permissionless blockchain where finality is measured in '
    'seconds [17].\n\n'
    'Academic research has meaningfully advanced detection capability through machine learning on transaction '
    'graphs. GraphSAGE and cognate graph neural architectures model wallet-level interactions to identify '
    'anomalies with considerable precision\u2014ROC-AUC scores exceeding 0.90 on benchmark datasets are now '
    'routine [5]. The limitation is consistent: these methods operate exclusively off-chain and enforce '
    'nothing on the ledger itself.\n\n'
    'Recent systematisations of knowledge underscore the migration towards blockchain-native RegTech. '
    'Mao et al. survey 41 commercial platforms and 28 academic prototypes spanning 2015\u20132025, constructing '
    'taxonomies for regulatory evolution, verification-layer compliance protocols, and lifecycle phases\u2014'
    'preventive, real-time, and investigative [4]. Their analysis confirms that Web3 architectures can '
    'facilitate transaction graph analysis, real-time risk scoring, cross-chain tracking, and '
    'privacy-preserving verification: capabilities that centralized systems struggle to replicate at '
    'equivalent fidelity. Gaps persist nonetheless. Cross-chain coverage remains incomplete, DeFi interaction '
    'monitoring is immature, privacy protocol oversight is nascent, and scalability constraints have not been '
    'resolved. The gulf between academic innovation and production deployment remains wide.\n\n'
    'Privacy-preserving methodologies occupy a central position in the balancing act between compliance '
    'obligations and data protection. Zero-knowledge proofs\u2014Groth16 circuits in particular\u2014permit '
    'verification of compliance attributes (non-sanctioned status, for instance) without exposing any '
    'underlying personal data [6]. This aligns naturally with GDPR constraints and lowers a material barrier '
    'to institutional DeFi adoption. Frameworks such as zkAML leverage ZK-SNARKs for whitelist-based AML '
    'within smart contracts, furnishing cryptographic proof of regulatory adherence whilst preserving user '
    'privacy. The principle extends to KYC/AML more broadly: verifiable credentials can attest to attributes '
    'without ever revealing raw data [16].\n\n'
    'Multi-agent and layered architectures present a more proactive paradigm. Khanvilkar et al. propose a '
    'decentralized multi-agent system for real-time compliance verification, distributing regulatory rules '
    'across agents via formal logic, consensus mechanisms, and zero-knowledge techniques [8]. Smart contracts '
    'supply immutable audit trails; experimental results report accuracy above 98% with low latency. The '
    'design eliminates single points of failure and accommodates cross-jurisdictional enforcement\u2014a '
    'non-trivial requirement for institutional participants operating across regulatory boundaries.\n\n'
    'Layered architectural patterns surface across adjacent domains. Compliance frameworks for distributed '
    'systems frequently adopt four-layer decompositions\u2014ingestion, ledger, processing, audit\u2014to guarantee '
    'consistency, immutability, and traceability. Gade describes one such ledger-centric, event-driven '
    'architecture for scalable, auditable systems [12]. Blockchain-enabled models in areas as diverse as '
    'human resource management and fixed-income tokenisation similarly employ layered structures to enforce '
    'transparency and codify rules [15]. These designs have inspired deterministic on-chain enforcement '
    'mechanisms, but they rarely integrate comprehensive AML pipelines with graph-temporal analytics.\n\n'
    'DeFi-specific scholarship has begun to examine institutional integration in earnest. Mohapatra and Raut '
    'investigate DeFi protocols for corporate treasury and liquidity management, recognising the efficiency '
    'gains from smart contracts and liquidity pools whilst emphasising the imperative of governance and '
    'regulatory alignment [11]. Olanrewaju addresses DeFi-based asset securitisation, cataloguing compliance '
    'and interoperability challenges and identifying oracle networks and KYC-integrated contracts as partial '
    'remedies [9].\n\n'
    'AMTTP builds squarely on these foundations whilst addressing their most conspicuous shortcoming: the '
    'absence of proactive enforcement. It couples off-chain analytics\u2014ensemble machine learning and temporal '
    'graph analysis\u2014with on-chain deterministic actions within a four-layer architecture. Where existing '
    'tools monitor and alert, AMTTP normalises heterogeneous signals into real-time compliance decisions: '
    'approve, review, escrow, or block. The integration of zkNAF for privacy-preserving attestations bridges '
    'the gap between reactive surveillance paradigms and the institutional-grade, pre-settlement enforcement '
    'that regulated participants require [2, 3].\n\n'
    'In sum, the extant literature demonstrates meaningful progress in detection and privacy but reveals a '
    'dearth of full-stack, production-grade protocols for deterministic AML enforcement tailored to '
    'institutional DeFi. AMTTP occupies precisely this gap, delivering an integrated, deployable system '
    'underpinned by open-source components [10, 13, 14].'
)

elevated['system_arch_intro'] = (
    'AMTTP is organised as a four-layer architecture, illustrated in Fig. 1. Each layer discharges a '
    'distinct set of responsibilities, spanning user interaction through to on-chain enforcement. This '
    'separation of concerns yields modularity, independent scalability, and long-term maintainability.'
)

elevated['layer1'] = (
    'Layer I provides the interfaces through which users, applications, and external systems interact with '
    'the protocol. It serves both human operators and programmatic consumers.\n\n'
    '**AMTTP Open-Source SDK:** The protocol furnishes client libraries in TypeScript and Python. These '
    'libraries manage transaction construction and signing via EIP-712 structured data [19], establish '
    'WebSocket connections for real-time event streaming, and expose batch scoring interfaces for '
    'high-throughput screening.\n\n'
    '**RESTful API Interface:** A FastAPI service on port 8000 exposes production endpoints for compliance '
    'evaluation. The /score endpoint assesses individual transactions and returns risk scores with '
    'accompanying explanations. The /batch endpoint processes multiple transactions in a single invocation. '
    'The /model/info endpoint surfaces metadata on the deployed machine learning model. The /alerts endpoint '
    'integrates with SIEM dashboards; /health furnishes readiness probes.\n\n'
    '**Web Application Layer:** Two complementary web applications address distinct constituencies. A '
    'Flutter-based consumer application integrates with MetaMask, enabling users to connect wallets and '
    'inspect risk scores prior to signing. A Next.js \u201cWar Room\u201d dashboard equips compliance officers with '
    'real-time alert visibility, case management workflows, risk score visualisations, and a zkNAF proof '
    'verification interface.'
)

elevated['layer2'] = (
    'Layer II houses the core decision-making engine. It ingests transaction data from Layer I, evaluates '
    'risk across multiple parallel services, and produces deterministic compliance decisions governed by a '
    'predefined rule matrix.\n\n'
    '**Compliance Orchestrator:** The Orchestrator (port 8007) functions as the central coordination point. '
    'Upon receipt of a transaction, it dispatches parallel requests to downstream services\u2014the ML Risk '
    'Engine, Graph API, Sanctions service, and Geo-risk service. It then aggregates responses, applies '
    'entity profile data, and synthesises a unified risk assessment.\n\n'
    '**Identity Management Module:** This module maintains profiles for known entities\u2014KYC status, '
    'historical transaction patterns, jurisdictional metadata\u2014and applies deterministic rules to incoming '
    'transactions.\n\n'
    '**Transaction Sequencing Module:** This module enriches transaction context with supplementary signals: '
    'sanctions matches, KYC status, rule-based alerts, geographic risk scores, fraud probability estimates, '
    'and confidence levels derived from the ML model.\n\n'
    '**ML Risk Engine (Production CPU):** The production risk engine operates on CPU infrastructure and loads '
    'model artefacts from a Docker volume. It applies the trained stacked ensemble\u2014GraphSAGE, LightGBM, '
    'and XGBoost\u2014to score incoming transactions.\n\n'
    '**Compliance Decision Matrix:** All signals are resolved through the deterministic action matrix '
    'presented in Table I.'
)

elevated['layer3'] = (
    'Layer III operates offline, dedicated to the development and validation of machine learning models '
    'consumed by the production risk engine. It executes on GPU infrastructure (Google Colab A100) and '
    'implements a multi-stage weak supervision methodology.\n\n'
    '**Stage 1\u2014Teacher Model Training (Hope Machine):** The pipeline commences with training data drawn from '
    'two sources: the BitcoinHeist and Ethereum Kaggle datasets. The combined corpus comprises 2,640,911 rows '
    'across 177 features, with a fraud prevalence of 0.87%\u2014a 113:1 class imbalance. Labels are independent '
    'and verified.\n\n'
    'VAE featurisation deploys a VAEWithAttention model on the A100 GPU. The variational autoencoder generates '
    'latent vectors that compress the original feature space whilst preserving salient information. '
    'Reconstruction error serves as an auxiliary anomaly signal, appended to the original 177 features.\n\n'
    '**Stage 2\u2014Weak Supervision Labelling:** AMTTP fuses three noisy label sources via consensus:'
)

elevated['layer3_bullets'] = [
    '**Label Source 1\u2014Teacher XGBoost (Weight: 0.4):** The teacher model generates predictions on 1,670 fresh transactions from BigQuery, utilising 171 engineered features.',
    '**Label Source 2\u2014Graph Rules (Weight: 0.3):** Grounded in FATF AML typologies, the system implements seven behavioural pattern detectors: SMURFING, FAN_OUT, FAN_IN, LAYERING, STRUCTURING, VELOCITY, and PEELING.',
    '**Label Source 3\u2014Memgraph Structural Analysis (Weight: 0.3):** Memgraph executes structural analysis computing degree centrality, community detection, and proximity to known illicit clusters.',
]

elevated['layer3_tail'] = (
    '**Stage 3a\u2014VAE/GNN Pipeline:** Using the weakly supervised labels, the pipeline trains a combined '
    'Variational Autoencoder and Graph Neural Network on Google Colab GPU infrastructure.\n\n'
    '**Stage 3b\u2014Production Models (CPU Deployed):** Trained model artefacts are mounted via Docker volume '
    'to the production ML Risk Engine in Layer II.'
)

elevated['layer4_intro'] = (
    'Layer IV underpins the entire protocol and bridges off-chain compliance decisions to deterministic '
    'on-chain enforcement.\n\n'
    '**API Gateway:** An NGINX reverse proxy manages TLS termination, rate limiting (100 requests per '
    'second), and path-based routing. A Cloudflare Tunnel provides secure external access.\n\n'
    '**On-Chain Verification\u2014Ethereum Smart Contracts:** AMTTP deploys 19 Solidity smart contracts on the '
    'Ethereum Sepolia testnet, enumerated in Table II.'
)

elevated['layer4_mid'] = (
    'The zkNAF circuits enable privacy-preserving compliance verification through three distinct proof '
    'types: KYC credential proofs, risk range proofs, and sanctions non-membership proofs.\n\n'
    '**Data Persistence Layer:** Multiple databases serve distinct storage requirements:'
)

elevated['layer4_bullets'] = [
    '**MongoDB** stores entity profiles, alerts, transaction histories, and audit trails.',
    '**Redis** provides in-memory caching for sessions, rate limits, and sender feature vectors.',
    '**Memgraph** maintains the graph database for entity relationships and risk propagation paths.',
    '**Helia/IPFS** stores immutable audit logs, evidence packages, and cryptographic proofs.',
    '**BigQuery** serves as the training data source with 30 days of Ethereum transactions (1.67 million transactions).',
]

elevated['layer4_tail'] = (
    '**Deployment and Observability:** The entire system is defined in docker-compose.full.yml, orchestrating '
    '17 containerised services. Prometheus and Grafana provide observability through dashboards monitoring '
    'service health, latency, model performance, and alert volumes.'
)

elevated['production_intro'] = (
    'AMTTP is realised as a containerised system deployed via Docker Compose. Table III catalogues the '
    '17 containerised services alongside their assigned ports and principal functions.'
)

elevated['security'] = (
    'AMTTP implements security controls at every stratum of the protocol architecture. The API Gateway '
    'provides TLS termination, rate limiting at 100 requests per second, and Cloudflare Tunnel integration. '
    'A dedicated UI Integrity Service monitors for front-end tampering attempts.\n\n'
    'The smart contract suite encompasses multiple security-focused components with multi-signature support '
    'via SafeModule integration, enabling M-of-N approval workflows for sensitive operations. CoreSecure and '
    'Streamlined variants furnish hardened and gas-optimised implementations fortified with reentrancy '
    'guards [18].\n\n'
    'The zkNAF framework supports three proof types, detailed in Table IV.'
)

elevated['conclusion'] = (
    'This paper has presented AMTTP Version 4.0: a four-layer architecture for deterministic compliance '
    'enforcement in institutional decentralized finance. The protocol addresses a structural deficiency in '
    'the existing DeFi stack\u2014the inability to intercept non-compliant transactions before settlement.\n\n'
    'The architecture integrates off-chain risk assessment with on-chain enforcement across four distinct '
    'layers. Layer I provides integration interfaces\u2014TypeScript and Python SDKs, REST APIs, and web '
    'applications. Layer II implements the core compliance logic, with an orchestrator coordinating ML risk '
    'scoring, graph analytics, and policy rules into a deterministic decision matrix. Layer III defines an '
    'offline machine learning pipeline that combines weak supervision from three independent label sources. '
    'Layer IV specifies the infrastructure: 19 smart contracts, an NGINX gateway, 17 containerised '
    'microservices, and a comprehensive data persistence layer.\n\n'
    'Three design elements warrant particular emphasis: a deterministic decision matrix guaranteeing that '
    'identical inputs invariably produce identical outcomes; the zkNAF framework for privacy-preserving '
    'compliance verification; and a multi-oracle architecture supporting configurable threshold signatures. '
    'AMTTP is specified as a production-ready system, with every component defined as infrastructure-as-code.'
)

elevated['future_work'] = (
    'Several extensions to the architecture are envisaged. Cross-chain expansion will extend the CrossChain '
    'contract to blockchain platforms beyond Ethereum Sepolia. Enhanced zero-knowledge capabilities\u2014'
    'additional proof types for more intricate compliance scenarios\u2014are under active investigation. Temporal '
    'graph analysis will be deployed to capture evolving transaction patterns over extended time horizons. '
    'Gas optimisation will further reduce on-chain operational costs. Deployment on Ethereum Layer 2 networks, '
    'notably Arbitrum and Optimism, will be evaluated for throughput and cost advantages. Support for '
    'regulatory frameworks beyond the UK FCA regime is planned, broadening the protocol\u2019s jurisdictional '
    'applicability.'
)

# =========================================================================
# CONSERVATIVE VERSION TEXT
# =========================================================================
conservative = {}

conservative['abstract'] = (
    'Decentralised finance (DeFi) has transformed financial markets by enabling automated transactions '
    'without intermediaries, yet institutional adoption remains limited due to the absence of enforceable '
    'compliance mechanisms. Unlike traditional finance, where transactions can be halted before settlement, '
    'DeFi transactions are irreversible once confirmed. Existing compliance tools focus on post-transaction '
    'monitoring rather than prevention, leaving regulated institutions exposed to regulatory risk. The UK '
    'Financial Conduct Authority has signalled through CP25/41 that firms must implement effective controls '
    'to prevent financial crime [1].\n\n'
    'This paper presents AMTTP Version 4.0, a four-layer architecture for deterministic compliance '
    'enforcement in institutional DeFi. Layer I provides SDKs, REST APIs, and web applications for user '
    'interaction. Layer II implements a compliance orchestrator coordinating machine learning risk scoring, '
    'graph analytics, sanctions screening, and policy rules into a deterministic decision matrix. Layer III '
    'defines an offline training pipeline using weak supervision from a teacher XGBoost model trained on '
    '2.64 million transactions, seven FATF AML pattern rules, and Memgraph structural analysis. Layer IV '
    'specifies the infrastructure, including 19 smart contracts on Ethereum Sepolia, 17 containerised '
    'microservices, and a data persistence layer comprising MongoDB, Redis, Memgraph, IPFS, and BigQuery.\n\n'
    'The security model incorporates multi-oracle threshold signatures, replay protection, and zkNAF '
    'zero-knowledge proofs for privacy-preserving verification of KYC credentials, risk range, and '
    'sanctions non-membership. Infrastructure security includes TLS encryption, rate limiting, Cloudflare '
    'Tunnel, and a UI Integrity Service.\n\n'
    'This paper demonstrates that deterministic compliance can be embedded into DeFi through a layered '
    'architecture. The complete specification is available as infrastructure-as-code.'
)

conservative['introduction'] = (
    'Decentralised Finance (DeFi) has transformed financial services by enabling open, automated '
    'transactions on blockchain networks without traditional intermediaries. Whilst this shift offers lower '
    'costs and faster processing, it creates significant compliance challenges, particularly for '
    'institutional users such as pension funds and asset managers who must adhere to strict regulatory '
    'standards.\n\n'
    'A fundamental problem is what regulators describe as the \u201ccompliance paradox\u201d in DeFi. Unlike '
    'traditional finance, where transactions can be stopped or reviewed before completion, DeFi transactions '
    'on blockchains such as Ethereum are immutable and final within seconds. This renders post-transaction '
    'monitoring ineffective and exposes institutions to regulatory fines, reputational harm, and interaction '
    'with illicit funds.\n\n'
    'Regulatory pressure is intensifying globally. In the United Kingdom, the Financial Conduct Authority '
    '(FCA) is implementing a comprehensive crypto-asset regime by 2027, requiring firms to meet Anti-Money '
    'Laundering (AML) and counter-terrorist financing standards, with full enforcement expected from '
    'October 2027 [2]. Globally, inconsistent regulations and strict capital requirements for crypto '
    'holdings further complicate institutional adoption.\n\n'
    'Existing AML tools for DeFi focus primarily on monitoring rather than prevention. Services such as '
    'Chainalysis and Elliptic provide off-chain analysis and alerts, but these require human review and '
    'cannot prevent high-risk transactions from settling. Academic advances such as GraphSAGE improve '
    'detection accuracy, yet remain separate from on-chain enforcement [5].\n\n'
    'This paper introduces the Anti-Money Laundering Transaction Trust Protocol (AMTTP), a system designed '
    'to embed compliance directly into the transaction process. AMTTP employs a four-layer architecture '
    'combining machine-learning-based risk scoring, graph analysis, and deterministic on-chain enforcement '
    'using smart contracts. The protocol also incorporates zkNAF, a zero-knowledge framework that enables '
    'privacy-preserving compliance verification whilst satisfying GDPR requirements [6]. The system is '
    'production-ready, deployed using a containerised stack and supported by an operational dashboard for '
    'real-time compliance monitoring.'
)

conservative['related_work'] = (
    'Research on compliance in Decentralised Finance (DeFi) has grown rapidly. It focuses on the challenges '
    'of applying traditional Anti-Money Laundering (AML) and Counter-Financing of Terrorism (CFT) rules to '
    'decentralised systems. Early studies examined general blockchain risks, whilst recent work addresses '
    'DeFi-specific issues such as privacy, real-time enforcement, and integration with institutional '
    'needs [1].\n\n'
    'Traditional AML tools rely on centralised monitoring and post-transaction analysis. Services such as '
    'Chainalysis and Elliptic use graph-based tracing and heuristics to detect suspicious patterns in '
    'cryptocurrency transactions. These tools provide alerts for human review but lack automated prevention '
    'in DeFi environments, where transactions settle irreversibly and quickly [17].\n\n'
    'Academic efforts have improved detection through machine learning on transaction graphs. For example, '
    'GraphSAGE and similar graph neural networks model wallet interactions to identify anomalies with high '
    'accuracy, such as ROC-AUC scores above 0.90 on benchmark datasets [5]. These methods, however, remain '
    'off-chain and do not enforce actions on the blockchain.\n\n'
    'Recent systematisations highlight the shift towards blockchain-native RegTech. Mao et al. present a '
    'comprehensive review of Web3 RegTech for Virtual Asset Service Provider (VASP) AML/CFT compliance [4]. '
    'They analyse 41 commercial platforms and 28 academic prototypes from 2015 to 2025. Their work develops '
    'taxonomies for regulatory evolution, compliance protocols across verification layers, and lifecycle '
    'phases (preventive, real-time, investigative). It demonstrates that Web3 solutions enable transaction '
    'graph analysis, real-time risk scoring, cross-chain tracking, and privacy-preserving verification\u2014'
    'features that are difficult to achieve in centralised systems. The study notes gaps in cross-chain '
    'coverage, DeFi interaction monitoring, privacy protocol oversight, and scalability, whilst emphasising '
    'the divide between academic innovations and industry deployment.\n\n'
    'Privacy-preserving techniques are essential for balancing compliance with data protection. Zero-knowledge '
    'proofs (ZKPs), including Groth16 circuits, allow verification of compliance attributes (e.g., '
    'non-sanctioned status) without revealing personal information [6]. This aligns with the GDPR and '
    'supports institutional DeFi adoption. Frameworks such as zkAML use ZK-SNARKs for whitelist-based AML in '
    'smart contracts, enabling cryptographic proof of regulatory adherence whilst preserving user privacy. '
    'Such approaches extend to KYC/AML in DeFi, where verifiable credentials prove attributes without '
    'exposing raw data [16].\n\n'
    'Multi-agent and layered architectures offer proactive enforcement. Khanvilkar et al. propose a '
    'decentralised multi-agent system for real-time compliance verification [8]. It distributes regulatory '
    'rules across agents using formal logic, consensus, and zero-knowledge techniques. Smart contracts '
    'provide immutable audit trails, achieving over 98% accuracy and low latency in experiments. This '
    'removes single points of failure and supports cross-jurisdictional enforcement.\n\n'
    'Layered designs appear in related domains. For compliance in distributed systems, frameworks often '
    'employ four layers (e.g., ingestion, ledger, processing, audit) to ensure consistency, immutability, '
    'and traceability. Gade describes a ledger-centric, event-driven architecture with four layers for '
    'scalable, auditable systems [12]. Similarly, blockchain-enabled models in other areas (e.g., HRM, '
    'fixed-income tokenisation) employ layered structures for transparency and rule enforcement [15]. These '
    'inspire deterministic on-chain mechanisms but rarely integrate full AML pipelines with graph-temporal '
    'analytics.\n\n'
    'DeFi-specific studies explore institutional integration. Mohapatra and Raut examine DeFi protocols for '
    'corporate treasury and liquidity management, highlighting smart contracts and liquidity pools for '
    'efficiency but stressing governance and regulatory alignment [11]. Olanrewaju discusses DeFi for asset '
    'securitisation, noting challenges in compliance and interoperability, with oracles and KYC-integrated '
    'contracts as partial solutions [9].\n\n'
    'AMTTP builds on these foundations. It addresses gaps in proactive enforcement by combining off-chain '
    'analytics (ensemble ML and temporal graphs) with on-chain deterministic actions in a four-layer '
    'architecture. Unlike monitoring-focused tools, AMTTP normalises signals for real-time decisions '
    '(approve/review/escrow/block). It incorporates zkNAF for privacy-preserving attestations, bridging '
    'reactive approaches and enabling institutional-grade compliance in DeFi without central '
    'intermediaries [2, 3].\n\n'
    'This related work demonstrates progress in detection and privacy but reveals limited full-stack, '
    'production-ready protocols for deterministic AML enforcement tailored to institutional DeFi. AMTTP '
    'advances this area by providing an integrated, deployable system with open-source components [10, 13, 14].'
)

conservative['system_arch_intro'] = (
    'AMTTP is organised as a four-layer architecture, as shown in Fig. 1. Each layer handles a distinct '
    'set of responsibilities, from user interaction to on-chain enforcement. This separation ensures '
    'modularity, scalability, and maintainability.'
)

conservative['layer1'] = (
    'Layer I provides the interfaces through which users, applications, and external systems interact with '
    'AMTTP. It is designed to be accessible to both human users and automated systems.\n\n'
    '**AMTTP Open-Source SDK:** The protocol offers client libraries in TypeScript and Python. These '
    'libraries handle transaction construction and signing using EIP-712 structured data [19], establish '
    'WebSocket connections for real-time event streaming, and provide batch scoring interfaces for '
    'high-throughput screening.\n\n'
    '**RESTful API Interface:** A FastAPI service running on port 8000 exposes production endpoints for '
    'compliance checking. The /score endpoint evaluates single transactions and returns risk scores with '
    'explanations. The /batch endpoint processes multiple transactions in one call. The /model/info endpoint '
    'provides metadata about the deployed machine learning model. The /alerts endpoint feeds into SIEM '
    'dashboards, and /health provides readiness probes.\n\n'
    '**Web Application Layer:** Two complementary web applications serve different user groups. A '
    'Flutter-based consumer application integrates with MetaMask, allowing users to connect their wallets '
    'and view risk scores before signing. A Next.js \u201cWar Room\u201d dashboard provides compliance officers with '
    'a real-time view of alerts, case management workflows, risk score visualisations, and a zkNAF proof '
    'verification interface.'
)

conservative['layer2'] = (
    'Layer II contains the core decision-making engine of AMTTP. It receives transaction data from Layer I, '
    'evaluates risk through multiple parallel services, and produces deterministic compliance decisions '
    'based on a predefined rule matrix.\n\n'
    '**Compliance Orchestrator:** The Orchestrator (port 8007) serves as the central coordination point. '
    'When a transaction arrives, it fans out parallel requests to downstream services including the ML Risk '
    'Engine, Graph API, Sanctions service, and Geo-risk service. It aggregates responses, applies entity '
    'profile data, and produces a unified risk assessment.\n\n'
    '**Identity Management Module:** This module maintains profiles for known entities, including KYC status, '
    'historical transaction patterns, and jurisdictional information. It applies deterministic rules to '
    'incoming transactions.\n\n'
    '**Transaction Sequencing Module:** This module enriches transaction context with additional signals '
    'including sanctions hits, KYC status, rule-based alerts, geographic risk scores, fraud probability '
    'estimates, and confidence levels from the ML model.\n\n'
    '**ML Risk Engine (Production CPU):** The production risk engine runs on CPU and loads model artefacts '
    'from a Docker volume. It applies the trained stacked ensemble (GraphSAGE + LightGBM + XGBoost) to '
    'score incoming transactions.\n\n'
    '**Compliance Decision Matrix:** All signals are resolved through a deterministic action matrix shown '
    'in Table I.'
)

conservative['layer3'] = (
    'Layer III operates offline to develop and validate the machine learning models used by the production '
    'risk engine. It runs on GPU infrastructure (Colab A100) and implements a multi-stage weak supervision '
    'approach.\n\n'
    '**Stage 1: Teacher Model Training (Hope Machine):** The pipeline begins with training data from two '
    'sources: the BitcoinHeist and Ethereum Kaggle datasets. The combined dataset contains 2,640,911 rows '
    'with 177 features. The fraud rate is 0.87%, representing a 113:1 class imbalance. Labels are '
    'independent and verified.\n\n'
    'VAE featurisation applies a VAEWithAttention model on the A100 GPU. The variational autoencoder '
    'generates latent vectors that compress the original features whilst preserving important information. '
    'Reconstruction error serves as an additional anomaly signal appended to the original 177 features.\n\n'
    '**Stage 2: Weak Supervision Labelling:** AMTTP combines three noisy label sources through consensus:'
)

conservative['layer3_bullets'] = [
    '**Label Source 1: Teacher XGBoost (Weight: 0.4):** The teacher model predicts on 1,670 fresh transactions from BigQuery using 171 engineered features.',
    '**Label Source 2: Graph Rules (Weight: 0.3):** Based on FATF AML typologies, the system implements seven behavioural patterns: SMURFING, FAN_OUT, FAN_IN, LAYERING, STRUCTURING, VELOCITY, and PEELING.',
    '**Label Source 3: Memgraph Structural Analysis (Weight: 0.3):** Memgraph performs structural analysis computing degree centrality, community detection, and proximity to known illicit clusters.',
]

conservative['layer3_tail'] = (
    '**Stage 3a: VAE/GNN Pipeline:** Using weakly supervised labels, the pipeline trains a combined '
    'Variational Autoencoder and Graph Neural Network model on Colab GPU.\n\n'
    '**Stage 3b: Production Models (CPU Deployed):** Trained model artefacts are mounted via Docker volume '
    'to the production ML Risk Engine in Layer II.'
)

conservative['layer4_intro'] = (
    'Layer IV provides the underlying infrastructure supporting the entire protocol and connects off-chain '
    'compliance decisions to deterministic on-chain enforcement.\n\n'
    '**API Gateway:** An NGINX reverse proxy handles TLS termination, rate limiting (100 requests per '
    'second), and path-based routing. A Cloudflare Tunnel provides secure external exposure.\n\n'
    '**On-Chain Verification: Ethereum Smart Contracts:** AMTTP deploys 19 Solidity smart contracts on the '
    'Ethereum Sepolia testnet, as shown in Table II.'
)

conservative['layer4_mid'] = (
    'The zkNAF circuits enable privacy-preserving compliance verification through three proof types: KYC '
    'credential proofs, risk range proofs, and sanctions non-membership proofs.\n\n'
    '**Data Persistence Layer:** Multiple databases serve distinct storage requirements:'
)

conservative['layer4_bullets'] = [
    '**MongoDB** stores entity profiles, alerts, transaction histories, and audit trails.',
    '**Redis** provides in-memory caching for sessions, rate limits, and sender feature vectors.',
    '**Memgraph** maintains the graph database for entity relationships and risk propagation paths.',
    '**Helia/IPFS** stores immutable audit logs, evidence packages, and cryptographic proofs.',
    '**BigQuery** serves as the training data source with 30 days of Ethereum transactions (1.67 million transactions).',
]

conservative['layer4_tail'] = (
    '**Deployment and Observability:** The entire system is defined in docker-compose.full.yml, orchestrating '
    '17 containerised services. Prometheus and Grafana provide observability through dashboards monitoring '
    'service health, latency, model performance, and alert volumes.'
)

conservative['production_intro'] = (
    'AMTTP is implemented as a containerised system deployed using Docker Compose. Table III lists the '
    '17 containerised services with their assigned ports and primary functions.'
)

conservative['security'] = (
    'AMTTP implements security controls across the protocol architecture. The API Gateway provides TLS '
    'termination, rate limiting at 100 requests per second, and Cloudflare Tunnel integration. A dedicated '
    'UI Integrity Service monitors for front-end manipulation attempts.\n\n'
    'The smart contract suite includes multiple security-focused components with multi-signature support '
    'through SafeModule integration, enabling M-of-N approval workflows for sensitive operations. CoreSecure '
    'and Streamlined variants provide hardened and gas-optimised implementations with reentrancy guards [18].\n\n'
    'The zkNAF framework enables three proof types as shown in Table IV.'
)

conservative['conclusion'] = (
    'This paper has presented AMTTP Version 4.0, a four-layer architecture for deterministic compliance '
    'enforcement in institutional decentralised finance. The protocol addresses a fundamental gap in existing '
    'DeFi infrastructure: the inability to prevent non-compliant transactions before settlement.\n\n'
    'The architecture integrates off-chain risk assessment with on-chain enforcement through four distinct '
    'layers. Layer I provides integration interfaces including TypeScript and Python SDKs, REST APIs, and '
    'web applications. Layer II implements core compliance logic with an orchestrator coordinating ML risk '
    'scoring, graph analytics, and policy rules into a deterministic decision matrix. Layer III defines an '
    'offline machine learning pipeline combining weak supervision from three label sources. Layer IV '
    'specifies the infrastructure, including 19 smart contracts, an NGINX gateway, 17 containerised '
    'microservices, and a comprehensive data persistence layer.\n\n'
    'The protocol incorporates design innovations including a deterministic decision matrix ensuring that '
    'identical inputs produce identical outcomes, the zkNAF framework for privacy-preserving compliance '
    'verification, and a multi-oracle architecture supporting configurable threshold signatures. AMTTP is '
    'specified as a production-ready system with all components defined as infrastructure-as-code.'
)

conservative['future_work'] = (
    'Several extensions to the AMTTP architecture are planned for future development. Cross-chain expansion '
    'will extend the CrossChain contract capability to additional blockchain platforms beyond Ethereum '
    'Sepolia. Enhanced zero-knowledge capabilities will explore additional proof types for more complex '
    'compliance scenarios. Temporal graph analysis will explore networks to capture evolving transaction '
    'patterns over time. Gas optimisation will further reduce costs for on-chain operations. Deployment '
    'options for Ethereum Layer 2 solutions such as Arbitrum and Optimism will be evaluated. Additional '
    'regulatory frameworks beyond UK FCA requirements will be supported.'
)

# Shared references
references = [
    'J. Joseph, F. John, and B. Teslim, "Anti-money laundering (AML) in DeFi platforms: Global enforcement trends," Multidisciplinary Research and Development Journal International, vol. 6, no. 1, pp. 1\u201312, 2025.',
    'Financial Conduct Authority, "CP25/41: Regulating Cryptoassets: Admissions & Disclosures," 2025.',
    'Financial Action Task Force, "International standards on combating money laundering and the financing of terrorism and proliferation: The FATF recommendations," 2023.',
    'Q. A. Mao, J. Wang, Y. Liu, L. Zhu, J. Chen, and J. Yan, "SoK: Web3 RegTech for cryptocurrency VASP AML/CFT compliance," arXiv preprint arXiv:2512.24888, 2025.',
    'W. Hamilton, Z. Ying, and J. Leskovec, "Inductive representation learning on large graphs," in Advances in Neural Information Processing Systems, 2017, pp. 1024\u20131034.',
    'J. Groth, "On the size of pairing-based non-interactive arguments," in Advances in Cryptology \u2013 EUROCRYPT 2016, Springer, 2016, pp. 305\u2013326.',
    'S. Herold, "Architectural compliance in component-based systems," Doctoral dissertation, Clausthal University of Technology, 2011.',
    'K. Khanvilkar, V. Shinde, and K. Kommuru, "Multi-agent collaboration for real-time compliance verification in decentralised fintech systems," in 2025 International Conference on Computer Communication and Informatics (ICCCI), IEEE, 2025, pp. 1\u20136.',
    'A. Olanrewaju, "Harnessing decentralised finance (DeFi) protocols for institutional asset securitisation in cross-jurisdictional banking ecosystems," International Journal of Science and Research Archive, vol. 15, no. 1, pp. 1119\u20131136, 2025.',
    'P. S. Sivaraju, "Driving operational excellence via multi-market network externalisation: A quantitative framework for optimising availability, security, and total cost in distributed systems," International Journal of Research and Applied Innovations, vol. 7, no. 5, pp. 11349\u201311365, 2024.',
    'P. Mohapatra and S. Raut, "The role of DeFi protocols in corporate treasury and liquidity management," Research Square, 2025, doi: 10.21203/rs.3.rs-7960323/v1.',
    'U. R. Gade, "Designing a ledger-centric, event-driven architecture for consistent and scalable systems," Journal of Engineering and Computer Sciences, vol. 4, no. 9, pp. 364\u2013371, 2025.',
    'F. J. Ogunmola and D. Ruby, "Automating the identification of interchange fee leakage in e-commerce: Using data models to spot when processors overcharge on interchange categories," 2025, unpublished.',
    'R. Campbell, "Synchronising concurrent security modernisation programmes: A systems integration framework for post-quantum cryptography, zero trust architecture, and AI security," 2026, unpublished.',
    'M. Madanchian and H. Taherdoost, "Blockchain-enabled human resource management for enhancing transparency, trust, and talent mobility in the digital era," Blockchains, vol. 4, no. 1, p. 2, 2026.',
    'A. S. Alego, R. Dellamatriz, V. Dario, E. Primo, J. B. Silveira, L. Miranda et al., "CRX: Implementing zero trust segmentation and AI in tokenised fixed-income platform for decentralised finance (DeFi)," Authorea Preprints, 2025.',
    'Memgraph, "Graph analytics for fraud detection," Technical White Paper, 2024.',
    'OpenZeppelin, "Contracts: Security," OpenZeppelin Documentation, 2024. [Online]. Available: https://docs.openzeppelin.com/contracts',
    'Ethereum Improvement Proposals, "EIP-712: Ethereum typed structured data hashing and signing," 2017. [Online]. Available: https://eips.ethereum.org/EIPS/eip-712',
]

elevated['references'] = references
conservative['references'] = references

# Generate both files
print("Generating elevated version DOCX...")
doc1 = make_ieee_docx('elevated', elevated)
doc1.save(r'C:\Users\Administrator\Downloads\AMTTP_Paper_Elevated.docx')
print("Saved: AMTTP_Paper_Elevated.docx")

print("Generating conservative version DOCX...")
doc2 = make_ieee_docx('conservative', conservative)
doc2.save(r'C:\Users\Administrator\Downloads\AMTTP_Paper_Conservative.docx')
print("Saved: AMTTP_Paper_Conservative.docx")

print("\nDone! Both files saved to Downloads.")

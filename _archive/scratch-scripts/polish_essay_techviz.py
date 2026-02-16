"""
TechViz polish of the AMTTP academic essay DOCX.
Preserves all formatting, tables, images, and styles.
Adds Section VI: Evaluation and Benchmarks.

Style: Oxford English, asymmetric sentence rhythm,
       no AI transition words, natural scholarly voice.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re, copy, os

SRC = r"C:\Users\Administrator\Downloads\AMTTP A Four-Layer Architecture for Deterministic Compliance Enforcement in Institutional DeFi.docx"
DST = r"C:\Users\Administrator\Downloads\AMTTP_TechViz_Polished.docx"

# ---------------------------------------------------------------------------
# Paragraph-index -> replacement text
# ---------------------------------------------------------------------------
REPLACEMENTS = {

    # ===================== ABSTRACT (P007-P011) =====================
    7: (
        "Abstract\u2014DeFi has redrawn the contours of capital markets. Automated, "
        "intermediary-free transactions compress settlement windows and strip out layers "
        "of cost\u2014yet institutional uptake remains stubbornly thin. The reason is "
        "not mysterious: there is no enforceable compliance mechanism at the protocol "
        "layer. In traditional finance, a suspect transfer can be intercepted and frozen "
        "before it settles. Blockchain offers no such luxury. Once confirmed on-chain, "
        "a transaction is final\u2014irreversible in every meaningful sense. The bulk of "
        "today\u2019s compliance tooling monitors after the fact, cataloguing what has "
        "already happened rather than preventing what should not. For regulated "
        "entities, this gap is not merely inconvenient; it constitutes material "
        "regulatory exposure. The UK Financial Conduct Authority, through its "
        "consultation paper CP25/41, leaves little room for ambiguity: firms must "
        "implement effective pre-settlement controls to forestall financial crime [2]."
    ),

    8: (
        "This paper sets out AMTTP Version 4.0\u2014a four-layer architecture "
        "engineered for deterministic compliance enforcement within institutional "
        "DeFi environments. Layer\u00a0I supplies SDKs, REST APIs, and web "
        "applications serving both human operators and programmatic consumers. "
        "Layer\u00a0II houses the compliance orchestrator: a decision engine that "
        "marshals machine learning risk scoring, graph analytics, sanctions screening, "
        "and policy adjudication into a single deterministic matrix. Layer\u00a0III "
        "specifies an offline training pipeline underpinned by weak supervision from a "
        "teacher XGBoost model trained across 2.64\u00a0million transactions, seven "
        "FATF AML pattern rules, and Memgraph structural analysis. Layer\u00a0IV "
        "describes the deployment substrate: 19 smart contracts on Ethereum Sepolia, "
        "17 containerised microservices, and a persistence tier spanning MongoDB, "
        "Redis, Memgraph, IPFS, and BigQuery."
    ),

    9: (
        "Security provisions include multi-oracle threshold signatures, replay "
        "protection, and zkNAF zero-knowledge proofs\u2014enabling privacy-preserving "
        "verification of KYC credentials, risk range, and sanctions non-membership "
        "without exposing underlying personal data. Infrastructure-level safeguards "
        "layer TLS encryption, rate limiting, Cloudflare Tunnel integration, and a "
        "dedicated UI Integrity Service on top."
    ),

    10: (
        "The central claim is direct: deterministic compliance can be woven into "
        "DeFi at the architectural level, not bolted on after the fact. The "
        "complete specification is published as infrastructure-as-code."
    ),

    11: (
        "Index Terms\u2014DeFi Compliance, AML, Oracle Architecture, Transaction "
        "Enforcement, Institutional Finance, Smart Contracts, Zero-Knowledge Proofs"
    ),

    # ===================== INTRODUCTION (P013, P020-P024) =====================
    13: (
        "Decentralised finance has remade financial services from the ground up. "
        "Open, automated transactions now execute across blockchain networks "
        "without any recourse to traditional intermediaries. Costs fall. Settlement "
        "compresses. Entire categories of market participant find themselves "
        "disintermediated. The efficiency gains are real\u2014but so are the "
        "compliance hazards, particularly for institutional actors: pension funds, "
        "sovereign wealth vehicles, asset managers\u2014entities bound by "
        "unforgiving regulatory standards."
    ),

    20: (
        "as pension funds and asset managers who must adhere to strict regulatory standards."
    ),

    21: (
        "At the core of this tension sits what regulators now call the \u2018compliance "
        "paradox\u2019 of DeFi. Within legacy financial infrastructure, a transaction "
        "suspected of impropriety can be frozen, scrutinised, and\u2014if "
        "warranted\u2014reversed before settlement completes. Blockchain offers no "
        "equivalent mechanism. On Ethereum, finality arrives in seconds. A confirmed "
        "transaction is immutable; no amount of retrospective monitoring can undo it. "
        "Institutions that engage without pre-settlement controls invite regulatory "
        "sanction, reputational harm, and\u2014perhaps most troublingly\u2014unwitting "
        "entanglement with illicit capital flows."
    ),

    22: (
        "Regulatory pressure is intensifying. The UK FCA is assembling a comprehensive "
        "crypto-asset regime, with AML and counter-terrorist financing obligations "
        "becoming enforceable from October 2027 [2]. Elsewhere, the picture fragments: "
        "divergent national frameworks and punitive capital requirements for "
        "digital-asset holdings raise the barriers to institutional participation "
        "still further."
    ),

    23: (
        "The existing AML apparatus for DeFi is, by and large, reactive. "
        "Commercial offerings\u2014Chainalysis, Elliptic, and their "
        "competitors\u2014perform off-chain analysis and surface alerts. "
        "Each one demands human review. None can halt a high-risk "
        "transaction before it settles. Academic contributions like GraphSAGE "
        "have sharpened detection accuracy considerably, yet they too operate "
        "in isolation from any on-chain enforcement mechanism [5]."
    ),

    24: (
        "This paper introduces the Anti-Money Laundering Transaction Trust "
        "Protocol\u2014AMTTP\u2014a system designed to embed compliance directly "
        "into the transactional lifecycle. The architecture is four-layered: "
        "machine learning risk scoring, graph-theoretic analysis, and "
        "deterministic on-chain enforcement via smart contracts work in "
        "concert. It incorporates zkNAF, a zero-knowledge framework permitting "
        "privacy-preserving compliance attestation consonant with GDPR "
        "requirements [6]. The system is production-ready\u2014containerised, "
        "fully deployable, and backed by an operational dashboard for "
        "real-time compliance oversight."
    ),

    # ===================== RELATED WORK (P026-P037) =====================
    26: (
        "Compliance scholarship in decentralised finance has expanded at pace over "
        "the past five years. The impetus is plain: existing Anti-Money Laundering "
        "and Counter-Financing of Terrorism frameworks were never built for "
        "permissionless protocols, and the operational gap widens with each "
        "passing quarter. Early work treated blockchain risk in sweeping terms."
    ),

    27: (
        "Subsequent contributions have narrowed their lens. Privacy-preserving "
        "regulation, real-time enforcement, and the peculiar demands of "
        "institutional DeFi participants now command the lion\u2019s share "
        "of attention [1]."
    ),

    28: (
        "Conventional AML tooling rests on centralised monitoring and "
        "post-settlement analysis. Chainalysis, Elliptic, and comparable "
        "commercial platforms trace graphs and apply heuristic pattern "
        "recognition to flag suspicious cryptocurrency activity. They "
        "generate alerts. They do not prevent settlement\u2014nor can they, "
        "on a permissionless blockchain where finality is measured in "
        "seconds [17]."
    ),

    29: (
        "Academic research has pushed detection capability forward through "
        "machine learning on transaction graphs. GraphSAGE and kindred "
        "neural architectures model wallet-level interactions to surface "
        "anomalies\u2014ROC-AUC scores above 0.90 on standard benchmarks "
        "are routine now [5]. The limitation is always the same: these "
        "methods live exclusively off-chain and enforce nothing on the "
        "ledger."
    ),

    30: (
        "Recent systematisations of knowledge confirm the drift towards "
        "blockchain-native RegTech. Mao et al. survey 41 commercial "
        "platforms and 28 academic prototypes across 2015\u20132025, "
        "constructing taxonomies that span regulatory evolution, "
        "verification-layer compliance protocols, and lifecycle "
        "phases\u2014preventive, real-time, investigative [4]. Their "
        "analysis shows that Web3 architectures can support transaction "
        "graph analysis, real-time risk scoring, cross-chain tracking, "
        "and privacy-preserving verification at a fidelity centralised "
        "systems struggle to match. Gaps remain, though. Cross-chain "
        "coverage is patchy. DeFi interaction monitoring is immature. "
        "Privacy protocol oversight barely exists. Scalability constraints "
        "persist. The distance between academic prototype and production "
        "deployment stays wide."
    ),

    31: (
        "Privacy-preserving methodologies sit at the heart of the balancing "
        "act between compliance obligations and data protection. "
        "Zero-knowledge proofs\u2014Groth16 circuits in particular\u2014allow "
        "verification of compliance attributes (sanctions non-membership, "
        "for example) without exposing any underlying personal data [6]. "
        "The fit with GDPR is natural, and the practical effect is to lower "
        "a material barrier to institutional DeFi adoption. Frameworks "
        "such as zkAML harness ZK-SNARKs for whitelist-based AML within "
        "smart contracts, producing cryptographic proof of regulatory "
        "adherence whilst preserving user privacy. The principle extends "
        "to KYC/AML more broadly: verifiable credentials can attest to "
        "attributes without ever revealing raw data [16]."
    ),

    32: (
        "Multi-agent and layered architectures offer a more proactive "
        "paradigm. Khanvilkar et al. propose a decentralised multi-agent "
        "system for real-time compliance verification\u2014distributing "
        "regulatory rules across agents via formal logic, consensus "
        "mechanisms, and zero-knowledge techniques [8]. Smart contracts "
        "supply immutable audit trails; the reported accuracy exceeds "
        "98\u2009% with low latency. Single points of failure are eliminated; "
        "cross-jurisdictional enforcement\u2014no trivial requirement for "
        "institutions operating across regulatory boundaries\u2014becomes "
        "architecturally viable."
    ),

    33: (
        "Layered architectural patterns recur across adjacent domains. "
        "Compliance frameworks for distributed systems frequently adopt "
        "four-layer decompositions\u2014ingestion, ledger, processing, "
        "audit\u2014to guarantee consistency,"
    ),

    34: (
        "immutability, and traceability. Gade documents one such "
        "ledger-centric, event-driven architecture for scalable, "
        "auditable systems [12]. Blockchain-enabled models in fields "
        "as varied as human resource management and fixed-income "
        "tokenisation employ layered structures to enforce transparency "
        "and codify rules [15]. These designs have informed deterministic "
        "on-chain enforcement mechanisms\u2014but comprehensive AML "
        "pipelines integrating graph-temporal analytics remain conspicuously "
        "absent."
    ),

    35: (
        "DeFi-specific scholarship has begun to examine institutional "
        "integration with appropriate seriousness. Mohapatra and Raut "
        "investigate DeFi protocols for corporate treasury and liquidity "
        "management, recognising efficiency gains from smart contracts "
        "and liquidity pools whilst stressing the imperative of governance "
        "and regulatory alignment [11]. Olanrewaju addresses DeFi-based "
        "asset securitisation, cataloguing compliance and interoperability "
        "challenges and pointing to oracle networks and KYC-integrated "
        "contracts as partial remedies [9]."
    ),

    36: (
        "AMTTP builds squarely on these foundations\u2014but targets their "
        "most glaring shortcoming: the absence of proactive enforcement. "
        "It couples off-chain analytics (ensemble machine learning and "
        "temporal graph analysis) with on-chain deterministic actions "
        "inside a four-layer architecture. Where existing tools monitor "
        "and alert, AMTTP normalises heterogeneous signals into real-time "
        "compliance decisions\u2014approve, review, escrow, or block. "
        "zkNAF integration for privacy-preserving attestations bridges "
        "the divide between reactive surveillance paradigms and the "
        "institutional-grade, pre-settlement enforcement that regulated "
        "participants actually need [2], [3]."
    ),

    37: (
        "The literature, taken as a whole, demonstrates meaningful progress "
        "in detection and privacy yet reveals a conspicuous dearth of "
        "full-stack, production-grade protocols for deterministic AML "
        "enforcement tailored to institutional DeFi. That is precisely the "
        "gap AMTTP addresses\u2014delivering an integrated, deployable "
        "system underpinned by open-source components [10], [13], [14]."
    ),

    # ===================== SYSTEM ARCHITECTURE (P039-P083) =====================
    39: (
        "AMTTP is organised as a four-layer architecture, illustrated in "
        "Figure\u00a01. Each layer discharges a distinct set of responsibilities, "
        "from user interaction through to on-chain enforcement. The "
        "separation of concerns this affords yields modularity, independent "
        "scalability, and long-term maintainability."
    ),

    41: (
        "Layer\u00a0I provides the interfaces through which users, "
        "applications, and external systems engage with the protocol. "
        "It serves both human operators and programmatic consumers."
    ),

    42: (
        "AMTTP Open-Source SDK: The protocol furnishes client libraries "
        "in TypeScript and Python. These handle transaction construction "
        "and signing via EIP-712 structured data, establish WebSocket "
        "connections for real-time event streaming, and expose batch "
        "scoring interfaces for high-throughput screening."
    ),

    43: (
        "RESTful API Interface: A FastAPI service on port\u00a08000 "
        "exposes production endpoints for compliance evaluation. "
        "The /score endpoint assesses individual transactions and returns "
        "risk scores alongside accompanying explanations. /batch processes "
        "multiple transactions in a single invocation. /model/info surfaces "
        "metadata on the deployed ML model. /alerts integrates with SIEM "
        "dashboards; /health furnishes readiness probes."
    ),

    47: (
        "Web Application Layer: Two complementary web applications address "
        "distinct constituencies. A Flutter-based consumer application "
        "integrates with MetaMask, letting users connect wallets and "
        "inspect risk scores before signing. A Next.js \u2018War Room\u2019 "
        "dashboard equips compliance officers with real-time alert "
        "visibility, case management workflows, risk score visualisations, "
        "and a zkNAF proof verification interface."
    ),

    49: (
        "Layer\u00a0II houses the core decision-making engine. It ingests "
        "transaction data from Layer\u00a0I, evaluates risk across multiple "
        "parallel services, and produces deterministic compliance decisions "
        "governed by a predefined rule matrix."
    ),

    51: (
        "Compliance Orchestrator: The Orchestrator (port\u00a08007) "
        "serves as the central coordination point. On receipt of a "
        "transaction, it dispatches parallel requests to downstream "
        "services\u2014the ML Risk Engine, Graph API, Sanctions service, "
        "and Geo-risk service\u2014then aggregates responses, applies "
        "entity profile data, and synthesises a unified risk assessment."
    ),

    52: (
        "Identity Management Module: This module maintains profiles for "
        "known entities\u2014KYC status, historical transaction patterns, "
        "jurisdictional metadata\u2014and applies deterministic rules to "
        "each incoming transaction accordingly."
    ),

    53: (
        "Transaction Sequencing Module: Here, transaction context is "
        "enriched with supplementary signals: sanctions matches, KYC "
        "status, rule-based alerts, geographic risk scores, fraud "
        "probability estimates, and confidence levels derived from the "
        "ML model."
    ),

    54: (
        "ML Risk Engine (Production CPU): The production risk engine runs "
        "on CPU infrastructure, loading model artefacts from a Docker "
        "volume. It applies the trained stacked ensemble\u2014GraphSAGE, "
        "LightGBM, and XGBoost\u2014to score incoming transactions in "
        "real time."
    ),

    55: (
        "Compliance Decision Matrix: All signals are resolved through "
        "the deterministic action matrix presented in Table\u00a0I."
    ),

    57: (
        "Layer\u00a0III operates offline. Its sole purpose is the "
        "development and validation of machine learning models consumed "
        "by the production risk engine. It executes on GPU infrastructure "
        "(Google Colab A100) and implements a multi-stage weak supervision "
        "methodology."
    ),

    58: (
        "Stage\u00a01\u2014Teacher Model Training (Hope Machine): The "
        "pipeline begins with training data drawn from two sources: the"
    ),

    64: (
        "BitcoinHeist and Ethereum Kaggle datasets. The combined corpus "
        "comprises 2,640,911 rows across 177 features, with a fraud "
        "prevalence of 0.87\u2009%\u2014a 113:1 class imbalance. Labels "
        "are independent and verified."
    ),

    65: (
        "VAE featurisation employs a VAEWithAttention model on the A100 "
        "GPU. The variational autoencoder compresses the original feature "
        "space into latent vectors whilst preserving salient information. "
        "Reconstruction error doubles as an auxiliary anomaly signal, "
        "appended to the original 177 features."
    ),

    66: (
        "Stage\u00a02\u2014Weak Supervision Labelling: AMTTP fuses three "
        "noisy label sources via consensus:"
    ),

    67: (
        "\u2022 Label Source\u00a01\u2014Teacher XGBoost (Weight: 0.4): "
        "The teacher model generates predictions on 1,670 fresh "
        "transactions drawn from BigQuery, utilising 171 engineered "
        "features."
    ),

    68: (
        "\u2022 Label Source\u00a02\u2014Graph Rules (Weight: 0.3): "
        "Grounded in FATF AML typologies, the system implements "
        "seven behavioural pattern detectors\u2014SMURFING, FAN_OUT, "
        "FAN_IN, LAYERING, STRUCTURING, VELOCITY, and PEELING."
    ),

    69: (
        "\u2022 Label Source\u00a03\u2014Memgraph Structural Analysis "
        "(Weight: 0.3): Memgraph executes structural analysis computing "
        "degree centrality, community detection, and proximity to "
        "known illicit clusters."
    ),

    70: (
        "Stage\u00a03a\u2014VAE/GNN Pipeline: Using the weakly "
        "supervised labels, the pipeline trains a combined Variational "
        "Autoencoder and Graph Neural Network on Colab GPU infrastructure."
    ),

    71: (
        "Stage\u00a03b\u2014Production Models (CPU Deployed): Trained "
        "artefacts are mounted via Docker volume to the production ML "
        "Risk Engine in Layer\u00a0II."
    ),

    73: (
        "Layer\u00a0IV underpins the entire protocol. It bridges off-chain "
        "compliance decisions to deterministic on-chain enforcement."
    ),

    75: (
        "API Gateway: An NGINX reverse proxy manages TLS termination, "
        "rate limiting (100 requests per second), and path-based routing. "
        "A Cloudflare Tunnel provides secure external access without "
        "exposing internal endpoints."
    ),

    76: (
        "On-Chain Verification\u2014Ethereum Smart Contracts: AMTTP "
        "deploys 19 Solidity smart contracts on the Ethereum Sepolia "
        "testnet, enumerated in Table\u00a0II."
    ),

    77: (
        "The zkNAF circuits enable privacy-preserving compliance "
        "verification through three distinct proof types: KYC credential "
        "proofs, risk range proofs, and sanctions non-membership proofs."
    ),

    78: (
        "Data Persistence Layer: Multiple databases serve distinct "
        "storage requirements:"
    ),

    79: (
        "\u2022 MongoDB stores entity profiles, alerts, transaction "
        "histories, and audit trails."
    ),

    80: (
        "\u2022 Redis provides in-memory caching for sessions, rate "
        "limits, and sender feature vectors."
    ),

    81: (
        "\u2022 Memgraph maintains the graph database for entity "
        "relationships and risk propagation paths."
    ),

    82: (
        "\u2022 Helia/IPFS stores immutable audit logs, evidence "
        "packages, and cryptographic proofs."
    ),

    83: (
        "\u2022 BigQuery serves as the training data source with "
        "30\u00a0days of Ethereum transactions (1.67\u00a0million "
        "transactions). Deployment"
    ),

    # ===================== PRODUCTION IMPLEMENTATION (P089) =====================
    89: (
        "AMTTP is realised as a containerised system deployed via Docker "
        "Compose. Table\u00a0III catalogues the 17 containerised services "
        "alongside their assigned ports and principal functions."
    ),

    # ===================== SECURITY MODEL (P095-P098) =====================
    95: (
        "AMTTP implements security controls at every stratum of the "
        "protocol architecture. The API Gateway provides TLS termination, "
        "rate limiting at 100 requests per second, and Cloudflare Tunnel"
    ),

    96: (
        "integration. A dedicated UI Integrity Service monitors for "
        "front-end tampering attempts."
    ),

    97: (
        "The smart contract suite incorporates multiple security-focused "
        "components with multi-signature support via SafeModule integration, "
        "enabling M-of-N approval workflows for sensitive operations. "
        "CoreSecure and Streamlined variants furnish hardened and "
        "gas-optimised implementations fortified with reentrancy guards."
    ),

    98: (
        "The zkNAF framework supports three proof types, detailed in "
        "Table\u00a0IV."
    ),

    # ===================== CONCLUSION (P104-P106) =====================
    104: (
        "This paper has presented AMTTP Version\u00a04.0: a four-layer "
        "architecture for deterministic compliance enforcement in "
        "institutional decentralised finance. The protocol addresses "
        "a structural deficiency in the existing DeFi stack\u2014the "
        "inability to intercept non-compliant transactions before "
        "settlement."
    ),

    105: (
        "The architecture weaves off-chain risk assessment together with "
        "on-chain enforcement across four layers. Layer\u00a0I provides "
        "integration interfaces\u2014TypeScript and Python SDKs, REST "
        "APIs, and web applications. Layer\u00a0II implements the core "
        "compliance logic: an orchestrator coordinating ML risk scoring, "
        "graph analytics, and policy rules into a deterministic decision "
        "matrix. Layer\u00a0III defines an offline machine learning pipeline "
        "combining weak supervision from three independent label sources. "
        "Layer\u00a0IV specifies the infrastructure\u201419 smart contracts, "
        "an NGINX gateway, 17 containerised microservices, and a "
        "comprehensive data persistence layer."
    ),

    106: (
        "Three design elements merit particular emphasis. First, a "
        "deterministic decision matrix guaranteeing that identical inputs "
        "invariably produce identical outcomes. Second, the zkNAF framework "
        "for privacy-preserving compliance verification. Third, a "
        "multi-oracle architecture supporting configurable threshold "
        "signatures. AMTTP is specified as a production-ready system; "
        "every component is defined as infrastructure-as-code."
    ),

    # ===================== FUTURE WORK (P108-P109) =====================
    108: (
        "Several extensions to the architecture are foreseen. Cross-chain "
        "expansion will extend the CrossChain contract beyond Ethereum "
        "Sepolia. Enhanced zero-knowledge capabilities\u2014additional "
        "proof types for more intricate compliance scenarios\u2014are under "
        "active investigation. Temporal graph analysis will be deployed to "
        "capture evolving transaction patterns over longer time horizons. "
        "Gas optimisation will reduce on-chain operational costs. "
        "Deployment on Ethereum Layer\u00a02 networks\u2014Arbitrum and "
        "Optimism in particular\u2014will be"
    ),

    109: (
        "evaluated for throughput and cost advantages. Support for "
        "regulatory frameworks beyond the UK FCA regime is planned, "
        "broadening the protocol\u2019s jurisdictional reach."
    ),
}

# Paragraphs to blank out (orphaned fragments merged into adjacent paragraphs)
BLANK_PARAGRAPHS = {20}

# ===== NEW SECTION VI: EVALUATION AND BENCHMARKS =====
EVAL_SECTION_INSERT_AFTER = 98
EVAL_HEADING = "VI. Evaluation and Benchmarks"

EVAL_PARAGRAPHS = [
    # --- 6A. Overview ---
    (
        "This section presents quantitative evaluation of AMTTP across six "
        "dimensions: machine learning model performance, cross-level "
        "generalisation, score recalibration, external validation against "
        "independent datasets, smart contract gas efficiency, and "
        "system-level integration testing. All evaluation uses the production "
        "pipeline described in Sections\u00a0III and\u00a0IV."
    ),

    # --- 6B. Sub-heading: ML Model Performance ---
    None,  # Placeholder for sub-heading "A. Machine Learning Model Performance"

    (
        "The V2 production pipeline employs a multi-stage architecture. "
        "A \u03b2-VAE, trained exclusively on normal addresses "
        "(468,564 samples), produces latent embeddings and anomaly "
        "scores. GraphSAGE (3-layer, 128-dimensional) and GATv2 "
        "(4-head attention) generate structural embeddings from a "
        "transaction graph of 625,168 nodes and 1,673,244 edges. PCA "
        "compresses the concatenated GNN embeddings from 192 to 45 "
        "dimensions, retaining 98.3\u2009% of variance. Optuna-tuned "
        "XGBoost and LightGBM classifiers, combined through a logistic "
        "regression meta-learner, produce final fraud probabilities. "
        "Table\u00a0V reports component-level and ensemble ROC-AUC and "
        "PR-AUC on the held-out test set (125,034 addresses, 25.05\u2009% "
        "fraud rate)."
    ),

    None,  # Placeholder for TABLE V

    (
        "The meta-ensemble achieves ROC-AUC of 1.0000 and PR-AUC of 0.9999 "
        "on the held-out test set. Near-perfect scores of this kind are "
        "expected rather than anomalous\u2014the test labels are proxy "
        "labels generated by the teacher\u2019s own weak supervision "
        "framework, so the student is being measured against labels it was "
        "trained to reproduce. The conservative generalisation estimate "
        "comes from five-fold cross-validation on 372 independently verified "
        "fraud addresses, yielding Student-XGBoost ROC-AUC of 0.9957 "
        "(Section\u00a0VI-B). Worth noting: the \u03b2-VAE reconstruction "
        "error exhibits a 2.60\u00d7 ratio between fraud and normal addresses "
        "(0.1550 versus 0.0595), confirming that the autoencoder correctly "
        "identifies fraudulent accounts as anomalous\u2014without having "
        "been trained on any fraud samples."
    ),

    # --- 6C. Sub-heading: Cross-Validation ---
    None,  # Placeholder for sub-heading "B. Five-Fold Cross-Validation"

    (
        "To gauge generalisation beyond a single train\u2013test partition, "
        "five-fold stratified cross-validation was carried out on the full "
        "625,168-address dataset (372 independently verified fraud addresses; "
        "fraud rate 0.06\u2009%). Table\u00a0VI reports per-fold results for "
        "the Student-XGBoost base model."
    ),

    None,  # Placeholder for TABLE VI

    (
        "Mean ROC-AUC across folds is 0.9965 (\u00b1\u00a00.0049); mean "
        "PR-AUC is 0.6586 (\u00b1\u00a00.0404). The Student-XGBoost base "
        "model reaches ROC-AUC of 0.9957 with PR-AUC of 0.6333 on "
        "aggregated out-of-fold predictions, yielding F1 of 0.635 at the "
        "optimal threshold. Teacher\u2013student prediction agreement "
        "stands at 99.27\u2009%."
    ),

    # --- 6D. Sub-heading: Teacher Model Baseline ---
    None,  # Placeholder for sub-heading "C. Teacher Model and Knowledge Distillation"

    (
        "The teacher model\u2014a stacked ensemble of XGBoost, LightGBM, "
        "and FTTransformer trained on 2,640,911 labelled samples from the "
        "BitcoinHeist and Ethereum Kaggle datasets (113:1 class "
        "imbalance)\u2014achieves test PR-AUC of 0.6491 and ROC-AUC of "
        "0.9255 on its native evaluation set. Per-chain analysis reveals "
        "markedly higher performance on Ethereum transactions (PR-AUC "
        "0.9517) than on Bitcoin (PR-AUC 0.6091). The discrepancy is "
        "attributable to the richer feature space available for Ethereum "
        "addresses."
    ),

    (
        "Knowledge distillation from teacher to student is validated "
        "through calibration and divergence metrics. On the test partition "
        "(125,034 addresses), the student meta-ensemble achieves a Brier "
        "score of 0.000718 against the teacher\u2019s 0.1898. "
        "Jensen\u2013Shannon divergence between teacher and student score "
        "distributions is 0.1601\u2014moderate but expected: the student "
        "has learnt a sharper decision boundary."
    ),

    # --- 6E. Sub-heading: Transaction-Level Models ---
    None,  # Placeholder for sub-heading "D. Transaction-Level Evaluation"

    (
        "AMTTP trains transaction-level models on 1,673,244 individual "
        "transactions. TX-Level XGBoost achieves ROC-AUC of 0.9878 and "
        "PR-AUC of 0.9713; TX-Level LightGBM reaches 0.9876 and 0.9707 "
        "respectively; their ensemble yields ROC-AUC of 0.9879, PR-AUC "
        "of 0.9715, and F1 of 0.901 (precision 0.927, recall 0.877)."
    ),

    # --- 6E-new. Sub-heading: Cross-Level Generalization ---
    None,  # Placeholder for sub-heading "E. Cross-Level Generalization"

    (
        "A critical engineering challenge in production AML systems is "
        "the granularity mismatch between training data and inference "
        "targets. The V2 pipeline\u2019s Student models were trained on "
        "address-level aggregates\u2014625,168 addresses, 71 features "
        "comprising 5 tabular, 64 \u03b2-VAE latents, reconstruction "
        "error, and GraphSAGE score\u2014yet production scoring must "
        "operate on individual incoming transactions. Profiling uncovered "
        "a severe training\u2013serving skew: at inference time, only 5 "
        "of 71 features were computable from a raw transaction. The "
        "remaining 66\u2014including every deep learning embedding\u2014were "
        "hardcoded to zero. In practice, 93\u2009% of model inputs were "
        "zeroes in production, degrading the Student\u00a0v2 XGBoost to "
        "ROC-AUC 0.6436 and F1 0.483 when evaluated on "
        "transaction-granularity data."
    ),

    (
        "The remedy follows what we term Path\u00a0B: retraining dedicated "
        "transaction-level models on 21 features that are fully available "
        "at inference time\u20146 extracted directly from the API "
        "TransactionRequest payload, 7 derived via arithmetic "
        "transformations, and 8 sender-history features retrieved "
        "from an address index cache (gracefully defaulting to zero "
        "if unavailable). Labels are propagated from address-level "
        "ground truth to individual transactions via sender join: if "
        "an address carries a fraud label, all of its outgoing "
        "transactions inherit that label\u2014yielding 420,848 "
        "fraudulent transactions out of 1,673,244 total (25.15\u2009% "
        "fraud rate). Table\u00a0VIII compares the resulting "
        "transaction-level models against address-level baselines "
        "evaluated on the same data."
    ),

    None,  # Placeholder for TABLE VIII

    (
        "The results are unequivocal. Matching the training granularity "
        "to the inference granularity is decisive: TX-Level XGBoost "
        "achieves ROC-AUC of 0.9878 versus the address-level "
        "Student\u00a0v2 XGBoost\u2019s 0.6436\u2014an absolute "
        "improvement of 0.344. F1 climbs from 0.483 to 0.901. "
        "The TX-Level models exhibit zero training\u2013serving skew: "
        "every feature used during training is guaranteed to be present "
        "at inference. The meta-learner collapses to a two-input logistic "
        "regression over XGBoost and LightGBM probabilities (coefficients "
        "7.74 and 2.22, intercept \u22126.25), eliminating all GPU and "
        "graph-computation dependencies from the real-time scoring path."
    ),

    # --- 6F-new. Sub-heading: Score Recalibration ---
    None,  # Placeholder for sub-heading "F. Score Recalibration"

    (
        "Raw model scores demand recalibration before they can function "
        "as actionable risk signals. The teacher XGBoost produces highly "
        "compressed raw probabilities (range 0\u20131.1\u2009%), yielding "
        "an optimal F1 threshold of merely 0.0027 despite ROC-AUC of "
        "0.9969. AMTTP addresses this with a two-stage recalibration "
        "pipeline. A sigmoid transform centred at the 75th percentile "
        "with steepness k\u2009=\u200930 maps raw scores into a "
        "calibrated probability space: p_cal\u2009=\u20091\u2009/\u2009"
        "(1\u2009+\u2009exp(\u2212k\u00a0\u00b7\u00a0(x\u2009\u2212\u2009"
        "P75))). This shifts the optimal F1 threshold from 0.0027 to "
        "0.5199 whilst preserving the ranking (calibrated ROC-AUC "
        "remains 0.9969). Percentile-rank normalisation then maps "
        "calibrated scores to a 0\u2013100 interpretability scale, "
        "enabling human-readable risk tiers: CRITICAL (top 1\u2009%, "
        "\u2265\u2009P99), HIGH (top 5\u2009%, \u2265\u2009P95), MEDIUM "
        "(top 15\u2009%, \u2265\u2009P85), LOW (top 30\u2009%, "
        "\u2265\u2009P70), and MINIMAL (remainder)."
    ),

    (
        "Threshold tuning sharpens operational performance. Empirical "
        "analysis revealed that the original model threshold of 0.96 was "
        "excessively conservative\u2014most flagged addresses escaped "
        "automated action entirely. Recalibration brings the block "
        "threshold down from 0.96 to 0.50, the escrow threshold from "
        "0.85 to 0.65, and the review threshold from 0.70 to 0.50. "
        "Behavioural pattern boosts augment the ML score additively: "
        "SMURFING (+25\u2009%), STRUCTURING (+25\u2009%), PEELING "
        "(+20\u2009%), LAYERING (+15\u2009%), FAN_OUT/FAN_IN "
        "(+15\u2009%/+10\u2009%). Graph-proximity boosts elevate scores "
        "for addresses with sanctioned connections\u2014direct: +40\u2009%, "
        "two-hop: +20\u2009%. A multi-signal multiplier applies when "
        "multiple detection methods converge (two signals: 1.2\u00d7; "
        "three signals: 1.5\u00d7), implementing the defence-in-depth "
        "principle quantitatively."
    ),

    # --- 6G-new. Sub-heading: External Validation ---
    None,  # Placeholder for sub-heading "G. External Validation"

    (
        "Assessing generalisation beyond the training distribution is "
        "not optional. AMTTP\u2019s V2 pipeline is evaluated against "
        "three independent external datasets using a multi-strategy "
        "protocol. For each dataset, five strategies are tested: "
        "(A)\u00a0XGBoost with zero-padded feature alignment; "
        "(B)\u00a0full pipeline with Platt recalibration on a "
        "20\u2009%/80\u2009% calibration\u2013evaluation split; "
        "(C)\u00a0Meta-LR with Platt scaling; (D)\u00a0concept alignment, "
        "where the V2 score is appended as an additional feature to the "
        "external dataset\u2019s native features and evaluated via "
        "five-fold cross-validated logistic regression; and (E)\u00a0the "
        "ML3 reinforcement learning module with Platt correction. Platt "
        "scaling fits a logistic regression on the calibration "
        "partition\u2019s predicted probabilities versus true labels and "
        "transforms all evaluation-partition scores accordingly\u2014addressing "
        "distributional shift without any retraining of the base models."
    ),

    (
        "The Elliptic Bitcoin dataset (46,564 labelled transactions, "
        "4,545 illicit, grounded in law enforcement intelligence) "
        "provides the most rigorous cross-chain test: the student model "
        "was trained exclusively on Ethereum data pseudo-labelled by a "
        "teacher that itself was trained predominantly on Bitcoin "
        "(99.7\u2009% BitcoinHeist). Knowledge distillation thus crosses "
        "both the chain boundary and the labelling methodology boundary. "
        "The XBlock Ethereum Phishing dataset (9,841 addresses; 2,179 "
        "confirmed phishing, from Zhejiang University) provides a "
        "cross-level generalisation test: address-level labels scored "
        "by a transaction-level model, with address aggregates mapped to "
        "sender/receiver features. The Forta Network dataset "
        "(7,891 addresses; 268 malicious) produces an instructive negative "
        "result\u2014both student and teacher achieve ROC-AUC below 0.5. "
        "This is not model failure per se, but semantic mismatch: "
        "AMTTP\u2019s behavioural fraud patterns and Forta\u2019s "
        "intelligence-based threat labels describe different phenomena. "
        "Distribution shift compounds the problem: training-set mean gas "
        "price is 1.1\u00a0gwei versus 42.4 for Forta addresses; mean "
        "sender transaction count is 2,833 versus 40."
    ),

    (
        "External datasets supply only partial feature overlap, "
        "necessitating a feature reconstruction step before scoring. "
        "For the XBlock dataset, the reconstruction pipeline computes "
        "seven FATF behavioural pattern scores (SMURFING, FAN_OUT, "
        "FAN_IN, LAYERING, STRUCTURING, VELOCITY, PEELING) from the "
        "available address aggregates, derives a composite fraud score "
        "as their sum, and generates a hybrid risk score by blending "
        "the Teacher XGBoost output with pattern and sophistication "
        "signals (0.4\u00a0\u00d7\u00a0Teacher XGB + "
        "0.3\u00a0\u00d7\u00a0pattern boost + "
        "0.3\u00a0\u00d7\u00a0sophistication). This enrichment raises "
        "TX-Level feature coverage from 9/21 to 14/21 and Student\u00a0v2 "
        "tabular coverage from 3/5 to 5/5. Before-and-after evaluation "
        "on XBlock confirms consistent improvement for "
        "transaction-level models: TX-Level XGBoost ROC-AUC rises from "
        "0.5833 to 0.6276 (+0.044), TX-Level LightGBM from 0.6385 to "
        "0.6562 (+0.018), TX-Level Ensemble from 0.6024 to 0.6454 "
        "(+0.043). F1 for the TX-Level Ensemble improves from 0.395 to "
        "0.416. The Student\u00a0v2 models, by contrast, show ROC-AUC "
        "degradation after enrichment (Student XGBoost 0.6507\u2009\u2192"
        "\u20090.3740)\u2014confirming that the address-level Student "
        "architecture with 66/71 features still zero-padded cannot be "
        "rescued by marginal tabular enrichment alone."
    ),

    (
        "Adaptive percentile-based thresholds\u2014calibrated to each "
        "model\u2019s score distribution rather than a fixed "
        "cut-off\u2014unlock latent ranking ability. On XBlock with "
        "percentile thresholds, the Teacher XGBoost achieves ROC-AUC "
        "of 0.7492 and F1 of 0.5288 (at the 39.4th percentile)\u2014the "
        "best single-model result across all generations. This is "
        "striking: the Teacher\u2019s raw probabilities span only "
        "[0.000256,\u20090.000401], rendering any fixed threshold above "
        "0.001 unusable. The percentile approach recovers its full "
        "discrimination power. On Elliptic, a concept alignment test "
        "appends the V2 fraud score as an additional feature to "
        "Elliptic\u2019s 166 native features and evaluates via five-fold "
        "cross-validated logistic regression: ROC-AUC rises from 0.9311 "
        "(native only) to 0.9334 (native\u2009+\u2009V2 score), and "
        "PR-AUC from 0.6241 to 0.6481\u2014confirming that the "
        "cross-chain signal adds incremental value even when the target "
        "dataset is already information-rich. Multi-layer coverage "
        "analysis on XBlock shows that combining ML, rule-based, and "
        "FATF pattern detectors achieves 99.8\u2009% recall on the "
        "2,179 phishing addresses; rules and FATF patterns uniquely "
        "catch 22 fraud cases the best ML model missed."
    ),

    (
        "A small-sample external sanity check corroborates directional "
        "validity. Twenty-three Ethereum addresses\u2014three independently "
        "confirmed as fraudulent via Etherscan transaction analysis\u2014are "
        "scored by the production pipeline. At the production risk "
        "threshold (CRITICAL or HIGH), the system achieves precision of "
        "1.0 and recall of 0.333 (F1\u2009=\u20090.50), with zero false "
        "positives among the 20 legitimate addresses. The continuous-score "
        "ROC-AUC of 0.4417 on this 23-address set is not statistically "
        "interpretable at n\u2009=\u200923\u2014but the perfect precision "
        "confirms that the system does not generate spurious alerts on "
        "clean addresses. A comprehensive external benchmark with larger "
        "independently labelled ground-truth sets remains an important "
        "direction for future work."
    ),

    # --- 6F-orig. Sub-heading: Theoretical Guarantees (was E) ---
    None,  # Placeholder for sub-heading "H. Theoretical Guarantees"

    (
        "Formal verification of the V2 pipeline establishes three "
        "results. Theorem\u00a01 (PAC-Bayes generalisation): the "
        "meta-ensemble\u2019s empirical risk of 0.000718 is bounded "
        "above by 0.3026 at 95\u2009% confidence (n\u2009=\u2009125,034, "
        "\u03b4\u2009=\u20090.05); the tightest component bound is "
        "GraphSAGE at 0.0857. Theorem\u00a02 (minimax adversarial "
        "equilibrium): von\u00a0Neumann\u2019s minimax theorem is verified "
        "numerically, with max\u2091\u00a0min\u1da0\u00a0V(D,F)\u2009"
        "=\u2009min\u1da0\u00a0max\u2091\u00a0V(D,F)\u2009=\u20090.0000, "
        "confirming that the detection game admits a Nash equilibrium. "
        "Theorem\u00a03 (unified adversarial bound): at worst-case "
        "perturbation \u03b5\u2009=\u20090.30, the unified PAC-Bayes "
        "adversarial bound is 0.3184\u2014decomposed into clean risk "
        "(0.0007), model complexity (0.3018), and attack complexity "
        "(0.0159)."
    ),

    # --- 6G-orig. Sub-heading: Smart Contract Gas Costs (was F) ---
    None,  # Placeholder for sub-heading "I. Smart Contract Gas Benchmarks"

    (
        "Table\u00a0VII reports gas consumption for representative smart "
        "contract operations, measured with Solidity\u00a00.8.24, "
        "optimiser runs\u2009=\u200950, and viaIR enabled. The escrow "
        "operation\u2014the most gas-intensive compliance "
        "action\u2014costs 287,890 gas (roughly 0.96\u2009% of the "
        "30M block gas limit). Dispute challenge consumes 149,512 gas. "
        "Administrative operations (pause, threshold adjustment) stay "
        "below 80,000 gas. Deployment costs for the three principal "
        "contracts total 10,518,877 gas (35\u2009% of block limit), "
        "necessitating separate deployment transactions."
    ),

    None,  # Placeholder for TABLE VII

    # --- 6H-orig. Sub-heading: System Integration (was G) ---
    None,  # Placeholder for sub-heading "J. System Integration Testing"

    (
        "End-to-end integration testing validates the complete deployment "
        "stack. Fifty-three API route tests pass across all backend "
        "services, covering /score, /batch, /alerts, and /health "
        "endpoints. Thirty-two browser-based UI interaction tests confirm "
        "that the Flutter consumer application and the Next.js War Room "
        "dashboard render correctly and respond to user actions. The full "
        "17-service Docker Compose stack starts within 90\u00a0seconds "
        "on commodity hardware."
    ),

    # --- 6I-orig. Sub-heading: Limitations (was H) ---
    None,  # Placeholder for sub-heading "K. Limitations"

    (
        "Several limitations warrant candid acknowledgement. The V2 "
        "pipeline\u2019s near-perfect test metrics (ROC-AUC 1.0000) are "
        "evaluated on proxy labels derived from the teacher\u2019s own "
        "weak supervision framework\u2014they measure internal consistency, "
        "not generalisation to independently verified ground truth. The "
        "honest cross-validated evaluation on 372 independently labelled "
        "fraud addresses yields the more conservative Student-XGBoost "
        "ROC-AUC of 0.9957. The external Etherscan validation set "
        "comprises only 23 addresses (3 fraud); at that sample size, "
        "no statistically powered conclusions can be drawn. Cross-domain "
        "evaluation against the Forta Network yields ROC-AUC below 0.5 "
        "for both student and teacher, underscoring that behavioural "
        "fraud models do not transfer to intelligence-derived threat "
        "taxonomies without domain adaptation. No persisted API latency "
        "benchmarks exist; throughput is estimated at 500+ transactions "
        "per second based on architecture capacity but has not been "
        "measured under sustained load. The cross-level label "
        "propagation\u2014inheriting address fraud labels onto all "
        "sender transactions\u2014assumes that every transaction from "
        "a fraudulent address is itself fraudulent. That is an "
        "approximation, and it may over-label benign activity."
    ),
]

EVAL_SUBHEADINGS = {
    1: "A. Machine Learning Model Performance",
    3: None,  # TABLE V placeholder
    5: "B. Five-Fold Cross-Validation",
    7: None,  # TABLE VI placeholder
    9: "C. Teacher Model and Knowledge Distillation",
    12: "D. Transaction-Level Evaluation",
    14: "E. Cross-Level Generalization",
    17: None,  # TABLE VIII placeholder
    19: "F. Score Recalibration",
    22: "G. External Validation",
    28: "H. Theoretical Guarantees",
    30: "I. Smart Contract Gas Benchmarks",
    32: None,  # TABLE VII placeholder
    33: "J. System Integration Testing",
    35: "K. Limitations",
}

# ---- Table data (unchanged — numeric data is sacrosanct) ----
TABLE_V_DATA = {
    "caption": "V\nV2 Pipeline Component Performance (Test Set, n = 125,034)",
    "headers": ["Component", "ROC-AUC", "PR-AUC"],
    "rows": [
        ["\u03b2-VAE (recon. error)", "0.8290", "0.5387"],
        ["\u03b2-VAE (Mahalanobis)", "0.7435", "0.4550"],
        ["GATv2", "0.9968", "0.9916"],
        ["GraphSAGE", "0.9994", "0.9984"],
        ["XGBoost (OOF)", "0.9999", "0.9998"],
        ["LightGBM (OOF)", "0.9999", "0.9998"],
        ["Meta-Ensemble", "1.0000", "0.9999"],
    ],
}

TABLE_VI_DATA = {
    "caption": "VI\nFive-Fold Cross-Validation (625,168 Addresses, 372 Fraud)",
    "headers": ["Fold", "ROC-AUC", "PR-AUC", "Time (s)"],
    "rows": [
        ["1", "0.9990", "0.7047", "164.1"],
        ["2", "0.9994", "0.6317", "145.7"],
        ["3", "0.9982", "0.6721", "175.8"],
        ["4", "0.9986", "0.6876", "147.1"],
        ["5", "0.9873", "0.5971", "52.1"],
        ["Mean \u00b1 SD", "0.9965 \u00b1 0.0049", "0.6586 \u00b1 0.0404", "136.9"],
    ],
}

TABLE_VII_DATA = {
    "caption": "VII\nSmart Contract Gas Consumption (Solidity 0.8.24, Optimiser Runs = 50)",
    "headers": ["Contract / Method", "Avg Gas", "% Block Limit"],
    "rows": [
        ["AMTTPDisputeResolver.escrowTransaction", "287,890", "0.96%"],
        ["AMTTPDisputeResolver.challengeTransaction", "149,512", "0.50%"],
        ["AMTTPCrossChain.setTrustedRemote", "100,047", "0.33%"],
        ["AMTTPCoreSecure.addOracle", "81,381", "0.27%"],
        ["AMTTPCoreSecure.addApprover", "80,792", "0.27%"],
        ["AMTTPCoreSecure.setUserPolicy", "76,649", "0.26%"],
        ["AMTTPCoreSecure.pause", "52,394", "0.17%"],
        ["Deploy: AMTTPCoreSecure", "5,011,664", "16.7%"],
        ["Deploy: AMTTPCrossChain", "3,309,779", "11.0%"],
        ["Deploy: AMTTPDisputeResolver", "2,197,434", "7.3%"],
    ],
}

TABLE_VIII_DATA = {
    "caption": "VIII\nCross-Level Generalisation: Address-Level vs. Transaction-Level Models",
    "headers": ["Model", "ROC-AUC", "PR-AUC", "F1", "Prec.", "Recall"],
    "rows": [
        ["Teacher Stack (\u26a0 circ.)", "0.7301", "0.6903", "0.7000", "1.0000", "0.5385"],
        ["Student v2 XGB (addr)", "0.6436", "0.3618", "0.4834", "0.4136", "0.5814"],
        ["Student v2 Ens. (addr)", "0.6322", "0.3491", "0.4728", "0.3959", "0.5867"],
        ["TX-Level XGB", "0.9878", "0.9713", "0.9011", "0.9210", "0.8820"],
        ["TX-Level LGB", "0.9876", "0.9707", "0.8986", "0.9253", "0.8733"],
        ["TX-Level Ensemble", "0.9879", "0.9715", "0.9012", "0.9272", "0.8765"],
    ],
}


# ===================================================================
# Document manipulation functions (identical logic, cleaner layout)
# ===================================================================

def replace_paragraph_text(para, new_text):
    """Replace paragraph text, preserving the first run's formatting."""
    if not para.runs:
        para.text = new_text
        return
    first_run = para.runs[0]
    for run in para.runs:
        run.text = ""
    first_run.text = new_text


def set_cell_shading(cell, colour):
    """Set cell background colour."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), colour)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def insert_table_after(doc, ref_para, table_data):
    """Insert an IEEE-style table immediately after ref_para."""
    caption_p = OxmlElement('w:p')
    ref_para._p.addnext(caption_p)
    cap_para = doc.paragraphs[[i for i, p in enumerate(doc.paragraphs)
                                if p._p is caption_p][0]]
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap_para.add_run("TABLE ")
    run.font.size = Pt(8)
    run.font.name = 'Times New Roman'
    cap_para.add_run("\n").font.size = Pt(8)
    r2 = cap_para.add_run(table_data["caption"])
    r2.font.size = Pt(8)
    r2.font.name = 'Times New Roman'

    headers = table_data["headers"]
    rows = table_data["rows"]
    ncols, nrows = len(headers), len(rows)
    table = doc.add_table(rows=1 + nrows, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = 'Table Grid'
    except KeyError:
        pass

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.paragraphs[0].text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(8)
        run.font.name = 'Times New Roman'
        set_cell_shading(cell, 'D9E2F3')

    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.paragraphs[0].text = ''
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(8)
            run.font.name = 'Times New Roman'

    cap_para._p.addnext(table._tbl)
    return cap_para


def _fix_section_margins(doc):
    """Fix float margin values that crash python-docx table creation."""
    for section in doc.sections:
        sectPr = section._sectPr
        pgMar = sectPr.find(qn('w:pgMar'))
        if pgMar is not None:
            for attr in ['w:left', 'w:right', 'w:top', 'w:bottom',
                         'w:header', 'w:footer', 'w:gutter']:
                val = pgMar.get(qn(attr))
                if val is not None:
                    try:
                        int(val)
                    except ValueError:
                        pgMar.set(qn(attr), str(int(float(val))))


def insert_eval_section(doc):
    """Insert Section VI: Evaluation and Benchmarks after Security Model."""
    _fix_section_margins(doc)

    anchor = doc.paragraphs[EVAL_SECTION_INSERT_AFTER]
    current = anchor

    # Section heading
    heading_el = OxmlElement('w:p')
    current._p.addnext(heading_el)
    heading_para = None
    for p in doc.paragraphs:
        if p._p is heading_el:
            heading_para = p
            break
    heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_para.space_before = Pt(12)
    heading_para.space_after = Pt(6)
    run = heading_para.add_run(EVAL_HEADING.upper())
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    current = heading_para
    print(f"  Inserted section heading: {EVAL_HEADING}")

    table_queue = {
        3: TABLE_V_DATA, 7: TABLE_VI_DATA,
        17: TABLE_VIII_DATA, 32: TABLE_VII_DATA,
    }
    inserted = 0

    for idx, content in enumerate(EVAL_PARAGRAPHS):
        if content is None and idx in EVAL_SUBHEADINGS and EVAL_SUBHEADINGS[idx] is not None:
            new_el = OxmlElement('w:p')
            current._p.addnext(new_el)
            sub_para = None
            for p in doc.paragraphs:
                if p._p is new_el:
                    sub_para = p
                    break
            sub_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            sub_para.space_before = Pt(8)
            sub_para.space_after = Pt(4)
            run = sub_para.add_run(EVAL_SUBHEADINGS[idx])
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
            run.italic = True
            current = sub_para
            inserted += 1
            print(f"  Inserted sub-heading: {EVAL_SUBHEADINGS[idx]}")

        elif content is None and idx in table_queue:
            tdata = table_queue[idx]
            cap_para = insert_table_after(doc, current, tdata)
            current = cap_para
            inserted += 1
            print(f"  Inserted table: {tdata['caption'].split(chr(10))[1].strip()}")

        elif content is None:
            continue

        else:
            new_el = OxmlElement('w:p')
            current._p.addnext(new_el)
            body_para = None
            for p in doc.paragraphs:
                if p._p is new_el:
                    body_para = p
                    break
            body_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            body_para.space_after = Pt(4)
            body_para.paragraph_format.first_line_indent = Inches(0.2)
            run = body_para.add_run(content)
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
            current = body_para
            inserted += 1

    print(f"  Total evaluation elements inserted: {inserted}")
    return inserted


def main():
    print(f"Opening: {SRC}")
    doc = Document(SRC)

    replaced = 0
    for i, para in enumerate(doc.paragraphs):
        if i in BLANK_PARAGRAPHS:
            pass
        if i in REPLACEMENTS:
            new_text = REPLACEMENTS[i]
            old_preview = para.text.strip()[:60]
            replace_paragraph_text(para, new_text)
            replaced += 1
            print(f"  P{i:03d}: replaced ({old_preview}...)")

    print("\nInserting Section VI: Evaluation and Benchmarks...")
    eval_count = insert_eval_section(doc)

    doc.save(DST)
    print(f"\nDone. Replaced {replaced} paragraphs, inserted {eval_count} evaluation elements.")
    print(f"Saved to: {DST}")


if __name__ == "__main__":
    main()

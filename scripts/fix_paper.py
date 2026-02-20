"""
Fix AMTTP paper: 
  1. Correct Table IV (ports, service names) to match actual deployment
  2. Fix contract count (18 -> 13) and service count (17 -> 24)
  3. Red-underline English issues (don't fix text, just mark them)
"""
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import copy

doc = Document(r'C:\Users\Administrator\Downloads\OLusegunAMTTP.docx')

# ═══════════════════════════════════════════════════════════════════════
# PART 1: FIX TABLE IV — Backend Services (para 155)
# ═══════════════════════════════════════════════════════════════════════
p155 = doc.paragraphs[155]
# Clear all existing runs
for run in p155.runs:
    run.text = ''

new_backend = (
    'NGINX Gateway 8888 Reverse proxy, TLS NGINX '
    'Flutter App 3010 End-user interface Flutter '
    'Next.js War Room 3006 Compliance dashboard Next.js '
    'Orchestrator 8007 Decision coord. FastAPI '
    'ML Risk Engine 8000 Risk scoring FastAPI '
    'Graph API 8001 Memgraph interface FastAPI '
    'FCA Compliance 8002 SAR, Travel Rule, FCA reports FastAPI '
    'Policy Service 8003 Policy CRUD, whitelist/blacklist FastAPI '
    'Sanctions Svc 8004 Watchlist screening FastAPI '
    'Monitoring Svc 8005 AML rules, alert management FastAPI '
    'Geo-Risk Svc 8006 Geographic risk FastAPI '
    'Integrity Svc 8008 UI tamper detection FastAPI '
    'Explainability 8009 ML XAI, typology analysis FastAPI '
    'zkNAF Demo 8010 Zero-knowledge proofs FastAPI '
    'Oracle Service 3001 KYC, PEP, EDD, bulk scoring Node.js '
    'Auth Gateway 8020 JWT authentication FastAPI '
)
p155.runs[0].text = new_backend
print('[FIXED] Table IV backend services — corrected ports, added missing services')

# ═══════════════════════════════════════════════════════════════════════
# PART 2: FIX TABLE IV — Infrastructure (para 156)
# Remove Prometheus/Grafana, add MinIO/Vault/Hardhat/Memgraph Lab
# ═══════════════════════════════════════════════════════════════════════
p156 = doc.paragraphs[156]
for run in p156.runs:
    run.text = ''

new_infra = (
    'MongoDB 27017 Document store MongoDB '
    'Redis 6379 In-memory cache Redis '
    'Memgraph 7687 Graph database Memgraph '
    'Memgraph Lab 3000 Graph visualisation Memgraph '
    'Helia/IPFS 5001 Immutable storage Helia '
    'MinIO 9000 S3-compatible object store MinIO '
    'HashiCorp Vault 8200 Secrets management Vault '
    'Hardhat Node 8545 Ethereum dev node Hardhat '
)
p156.runs[0].text = new_infra
print('[FIXED] Table IV infrastructure — removed Prometheus/Grafana, added real services')

# ═══════════════════════════════════════════════════════════════════════
# PART 3: FIX CONTRACT COUNT (18 -> 13)
# ═══════════════════════════════════════════════════════════════════════
# Para 7: abstract
p7 = doc.paragraphs[7]
for run in p7.runs:
    run.text = run.text.replace('18 smart contracts', '13 smart contracts')
print('[FIXED] Abstract: 18 -> 13 smart contracts')

# Para 140: "Eigh teen Solidity smart contracts"
p140 = doc.paragraphs[140]
for run in p140.runs:
    run.text = run.text.replace('Eigh teen', 'Thirteen')
    run.text = run.text.replace('Eighteen', 'Thirteen')
    run.text = run.text.replace('eighteen', 'thirteen')
print('[FIXED] Section III-D: Eighteen -> Thirteen contracts')

# Para 277: conclusion "18 smart contracts"  
p277 = doc.paragraphs[277]
for run in p277.runs:
    run.text = run.text.replace('18 smart contracts', '13 smart contracts')
    run.text = run.text.replace('18 Smart Contracts', '13 Smart Contracts')
print('[FIXED] Conclusion: 18 -> 13 smart contracts')

# ═══════════════════════════════════════════════════════════════════════
# PART 4: FIX SERVICE COUNT (17 -> 24)
# ═══════════════════════════════════════════════════════════════════════
p149 = doc.paragraphs[149]
for run in p149.runs:
    run.text = run.text.replace('17 services', '24 services')
print('[FIXED] Section IV: 17 -> 24 services')

# ═══════════════════════════════════════════════════════════════════════
# PART 5: RED UNDERLINE English issues (mark only, don't fix)
# ═══════════════════════════════════════════════════════════════════════

def red_underline_phrase(paragraph, phrase):
    """Find phrase in paragraph runs and apply red underline formatting."""
    # Build full text from runs with positions
    full_text = ''
    run_map = []  # (run_index, char_start_in_run, char_end_in_run)
    for ri, run in enumerate(paragraph.runs):
        for ci, ch in enumerate(run.text):
            run_map.append((ri, ci))
        full_text += run.text
    
    # Find phrase in full text
    idx = full_text.find(phrase)
    if idx == -1:
        # Try case-insensitive
        idx = full_text.lower().find(phrase.lower())
    if idx == -1:
        return False
    
    phrase_end = idx + len(phrase)
    
    # Identify which runs contain the phrase
    affected_runs = set()
    for pos in range(idx, min(phrase_end, len(run_map))):
        affected_runs.add(run_map[pos][0])
    
    # We need to split runs so the phrase is isolated
    # Strategy: rebuild the paragraph's runs
    runs_data = []
    for run in paragraph.runs:
        try:
            fs = run.font.size
        except (ValueError, TypeError):
            fs = None
        runs_data.append({
            'text': run.text,
            'bold': run.bold,
            'italic': run.italic,
            'font_name': run.font.name,
            'font_size': fs,
        })
    
    # Find positions in the concatenated text
    new_runs = []
    pos = 0
    for rd in runs_data:
        run_text = rd['text']
        run_start = pos
        run_end = pos + len(run_text)
        
        # Check overlap with phrase
        overlap_start = max(run_start, idx)
        overlap_end = min(run_end, phrase_end)
        
        if overlap_start < overlap_end:
            # This run contains part of the phrase
            # Split into: before | phrase_part | after
            before = run_text[:overlap_start - run_start]
            phrase_part = run_text[overlap_start - run_start:overlap_end - run_start]
            after = run_text[overlap_end - run_start:]
            
            if before:
                new_runs.append({**rd, 'text': before, 'red_underline': False})
            if phrase_part:
                new_runs.append({**rd, 'text': phrase_part, 'red_underline': True})
            if after:
                new_runs.append({**rd, 'text': after, 'red_underline': False})
        else:
            new_runs.append({**rd, 'text': run_text, 'red_underline': False})
        
        pos = run_end
    
    # Clear existing runs
    for run in paragraph.runs:
        run.text = ''
    
    # Remove all but first run from XML
    p_elem = paragraph._element
    runs_xml = p_elem.findall(qn('w:r'))
    for rx in runs_xml[1:]:
        p_elem.remove(rx)
    
    # Now rebuild
    first = True
    for nr in new_runs:
        if first and runs_xml:
            run = paragraph.runs[0]
            first = False
        else:
            run = paragraph.add_run()
        
        run.text = nr['text']
        if nr.get('bold'):
            run.bold = True
        if nr.get('italic'):
            run.italic = True
        if nr.get('font_name'):
            run.font.name = nr['font_name']
        if nr.get('font_size'):
            run.font.size = nr['font_size']
        
        if nr.get('red_underline'):
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            run.underline = True
    
    return True

# English issues to mark with red underline
english_issues = [
    (83,  'fans out'),
    (274, 'Candour about shortcomings matters'),
    (297, 'artefacts are available'),
]

# Also search more broadly for issues that may be in different paragraphs
broad_issues = [
    'multi-agents are being an evolution',
    'Its responsible',
    'data speaks for itself',
    'tax of determinism',
    'highest priority engineering task',
    'needs to be done right away',
    'Ruby and XGBoost',
    'AQ1',
    'recalcitration',
    'no single person has control',
    'succeeded to produce',
    'baked in',
    'so-called',
]

# First do the known locations
for para_idx, phrase in english_issues:
    p = doc.paragraphs[para_idx]
    if red_underline_phrase(p, phrase):
        print(f'[UNDERLINED] Para {para_idx}: "{phrase}"')
    else:
        print(f'[NOT FOUND]  Para {para_idx}: "{phrase}"')

# Then search all paragraphs for the broader patterns
for phrase in broad_issues:
    for i, p in enumerate(doc.paragraphs):
        if phrase.lower() in p.text.lower():
            if red_underline_phrase(p, phrase):
                print(f'[UNDERLINED] Para {i}: "{phrase}"')
            break  # Only first occurrence

# ═══════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════
output_path = r'C:\Users\Administrator\Downloads\OLusegunAMTTPnewedited_FIXED.docx'
doc.save(output_path)
print(f'\nSaved to: {output_path}')
print('Done! Architecture/table FIXED. English issues RED-UNDERLINED (not changed).')

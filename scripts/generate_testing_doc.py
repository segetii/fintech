"""
Generate AMTTP Testing Guide as Word Document
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

# Create document
doc = Document()

# Title
title = doc.add_heading('AMTTP Testing Guide', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('Version 2.3 - January 22, 2026')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Section 1: Application URLs
doc.add_heading('📋 Application URLs', level=1)

url_table = doc.add_table(rows=6, cols=3)
url_table.style = 'Table Grid'

# Header row
url_table.rows[0].cells[0].text = 'Application'
url_table.rows[0].cells[1].text = 'URL'
url_table.rows[0].cells[2].text = 'Description'

# Data rows
urls = [
    ('Flutter App', 'http://localhost:3010', 'Main mobile/web app'),
    ('Next.js App', 'http://localhost:3006', 'Web dashboard'),
    ('Risk Engine API', 'http://localhost:8002', 'ML scoring service'),
    ('Memgraph Lab', 'http://localhost:3000', 'Graph database UI'),
    ('Memgraph Bolt', 'bolt://localhost:7687', 'Graph database connection'),
]

for i, (app, url, desc) in enumerate(urls, 1):
    url_table.rows[i].cells[0].text = app
    url_table.rows[i].cells[1].text = url
    url_table.rows[i].cells[2].text = desc

doc.add_paragraph()

# Section 2: Demo Login Credentials
doc.add_heading('🔐 Demo Login Credentials', level=1)

doc.add_heading('Focus Mode Users (R1-R2) - Simplified Interface', level=2)

focus_table = doc.add_table(rows=3, cols=5)
focus_table.style = 'Table Grid'

focus_table.rows[0].cells[0].text = 'Role'
focus_table.rows[0].cells[1].text = 'Email'
focus_table.rows[0].cells[2].text = 'Password'
focus_table.rows[0].cells[3].text = 'Mode'
focus_table.rows[0].cells[4].text = 'Capabilities'

focus_users = [
    ('R1 - End User', 'user@amttp.io', 'user123', 'Focus Mode', 'Personal wallet, send/receive'),
    ('R2 - End User (PEP)', 'pep@amttp.io', 'pep123', 'Focus Mode', 'Same as R1 + enhanced monitoring'),
]

for i, row_data in enumerate(focus_users, 1):
    for j, cell_data in enumerate(row_data):
        focus_table.rows[i].cells[j].text = cell_data

doc.add_paragraph()

doc.add_heading('War Room Users (R3-R6) - Full Analytics Dashboard', level=2)

war_table = doc.add_table(rows=5, cols=5)
war_table.style = 'Table Grid'

war_table.rows[0].cells[0].text = 'Role'
war_table.rows[0].cells[1].text = 'Email'
war_table.rows[0].cells[2].text = 'Password'
war_table.rows[0].cells[3].text = 'Mode'
war_table.rows[0].cells[4].text = 'Capabilities'

war_users = [
    ('R3 - Institution Ops', 'ops@amttp.io', 'ops123', 'War Room (View)', 'Read-only analytics'),
    ('R4 - Compliance Officer', 'compliance@amttp.io', 'comply123', 'War Room (Full)', 'Policy management, multisig'),
    ('R5 - Platform Admin', 'admin@amttp.io', 'admin123', 'War Room (Admin)', 'User management, config'),
    ('R6 - Super Admin', 'super@amttp.io', 'super123', 'War Room (Super)', 'Emergency override ONLY'),
]

for i, row_data in enumerate(war_users, 1):
    for j, cell_data in enumerate(row_data):
        war_table.rows[i].cells[j].text = cell_data

doc.add_paragraph()

# Section 3: Test Cases by Role
doc.add_heading('🧪 Test Cases by Role', level=1)

doc.add_heading('R1/R2 - Focus Mode Tests', level=2)
doc.add_paragraph('URL: http://localhost:3010')

focus_tests = [
    'Authentication',
    '  ☐ Login with email/password',
    '  ☐ Login via demo account tile',
    '  ☐ Logout functionality',
    '  ☐ Session persistence after refresh',
    '',
    'Home Page',
    '  ☐ Dashboard loads correctly',
    '  ☐ Balance display',
    '  ☐ Recent transactions list',
    '',
    'Wallet',
    '  ☐ View wallet address',
    '  ☐ Copy address to clipboard',
    '',
    'Transfer',
    '  ☐ Send transaction form',
    '  ☐ Pre-transaction trust check',
    '',
    'History',
    '  ☐ Transaction list loads',
    '  ☐ Transaction details view',
]

for test in focus_tests:
    doc.add_paragraph(test)

doc.add_heading('R3 - Institution Ops Tests', level=2)

r3_tests = [
    'War Room Landing',
    '  ☐ Dashboard overview loads',
    '  ☐ Statistics cards display',
    '',
    'Flagged Queue',
    '  ☐ List of flagged transactions',
    '  ☐ Cannot approve/reject (view only)',
    '',
    'Graph Explorer',
    '  ☐ Graph visualization loads',
    '  ☐ Node click shows details',
]

for test in r3_tests:
    doc.add_paragraph(test)

doc.add_heading('R4 - Compliance Officer Tests', level=2)
doc.add_paragraph('All R3 capabilities PLUS:')

r4_tests = [
    '  ☐ Approve transaction',
    '  ☐ Reject transaction',
    '  ☐ Policy management',
    '  ☐ Generate compliance report',
]

for test in r4_tests:
    doc.add_paragraph(test)

doc.add_heading('R5 - Platform Admin Tests', level=2)
doc.add_paragraph('All R4 capabilities PLUS:')

r5_tests = [
    '  ☐ View all users',
    '  ☐ Create/edit user',
    '  ☐ System configuration',
    '  ☐ View audit logs',
]

for test in r5_tests:
    doc.add_paragraph(test)

doc.add_heading('R6 - Super Admin Tests', level=2)

r6_tests = [
    '  ☐ Emergency pause system',
    '  ☐ Override frozen transaction',
    '  ☐ No access to Detection Studio',
]

for test in r6_tests:
    doc.add_paragraph(test)

doc.add_paragraph()

# Section 4: Data Loaded
doc.add_heading('📊 Data Loaded', level=1)

data_info = doc.add_table(rows=4, cols=2)
data_info.style = 'Table Grid'

data_rows = [
    ('Address Nodes', '2,830'),
    ('Transaction Edges', '5,000'),
    ('Sanctioned Addresses', '2'),
    ('Data Source', 'eth_transactions_full_labeled.parquet (30% sample)'),
]

for i, (label, value) in enumerate(data_rows):
    data_info.rows[i].cells[0].text = label
    data_info.rows[i].cells[1].text = value

doc.add_paragraph()

# Section 5: Test Execution Log
doc.add_heading('📝 Test Execution Log', level=1)

log_table = doc.add_table(rows=7, cols=4)
log_table.style = 'Table Grid'

log_table.rows[0].cells[0].text = 'Test Case'
log_table.rows[0].cells[1].text = 'Status'
log_table.rows[0].cells[2].text = 'Date'
log_table.rows[0].cells[3].text = 'Notes'

log_rows = [
    ('Flutter App Response', '✅ PASS', '2026-01-22', 'HTML loads at :3010'),
    ('Next.js App Response', '✅ PASS', '2026-01-22', 'Server running at :3006'),
    ('Risk Engine Health', '✅ PASS', '2026-01-22', 'Model loaded, healthy'),
    ('Memgraph Connection', '✅ PASS', '2026-01-22', '2,830 nodes, 5,000 edges'),
    ('Docker Containers', '✅ PASS', '2026-01-22', '12 containers running'),
    ('Data Ingestion', '✅ PASS', '2026-01-22', '5,000 transactions loaded'),
]

for i, row_data in enumerate(log_rows, 1):
    for j, cell_data in enumerate(row_data):
        log_table.rows[i].cells[j].text = cell_data

doc.add_paragraph()

# Footer
footer = doc.add_paragraph()
footer.add_run(f'Document generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save document
doc.save(r'C:\amttp\AMTTP_Testing_Guide.docx')
print('Document saved to C:\\amttp\\AMTTP_Testing_Guide.docx')
